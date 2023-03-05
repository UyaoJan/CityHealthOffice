from tkinter import *
from PIL import ImageTk,Image
from tkcalendar import Calendar,DateEntry
import calendar
from tkinter import ttk
from tkinter import filedialog
from PIL import Image, ImageTk
from tkinter import messagebox
import datetime
from datetime import date, datetime
import employee, LoginPage
import docx
from pathlib import Path
from docxtpl import DocxTemplate
from docxtpl import InlineImage
from docx.shared import Inches
import numpy as np
import matplotlib.pyplot as plt
import win32api, os, time
import win32print
import summary_filter

from dotenv import load_dotenv

# get the handle to the taskbar
# h = windll.user32.FindWindowA(b'Shell_TrayWnd', None)


class Main:
    def __init__(self,init):
        env_loc="config.env"
        load_dotenv(env_loc)
        self.user=init
        self.Dashboard_GUI = None
        self.Page_Summary = None
        self.Page_FrontDesk = None
        self.Page_XRAY = None
        global PageOpen
        PageOpen = 1
        self.Value_Laboratory = ["Laboratory","X_RAY"]
        self.Close_ID=["Laboratory","X_RAY","Summary","FrontDesk"]
        
        # windll.user32.ShowWindow(h, 0)
        # self.client_bmonth=None

        self.test_date=date.today()
    
        #Checkbox Def
    
    def HOME_PAGE(self,Close):
        width= 1000
        height= 500
        self.Dashboard_GUI.resizable(False,False)
        self.Dashboard_GUI.attributes("-fullscreen", False)
        self.Dashboard_GUI.geometry(f'{width}x{height}+{180}+{80}')
        self.Page_Dashboard.pack(expand=1, fill=BOTH)
        
        if Close == "Laboratory":
            self.Page_Laboratory.destroy()
        elif Close == "X_RAY":
            self.Page_XRAY.destroy()
        elif Close == "Summary":
            self.Page_Summary.destroy()
        elif Close == "FrontDesk":
            self.Page_FrontDesk.destroy()

    def showCheckbox(self):
        errors=0
        chosen_serve=[]
        if Name_Entry.get() is None:
            errors +=1
        name=Name_Entry.get()
        if AGE_Entry.get() is None or AGE_Entry.get().isnumeric()==False:
            errors+=1
        age=AGE_Entry.get()

        bdate=str(Year_Birth.get())+'-'+str(self.client_bmonth)+'-'+str(Day_Birth.get())
        bdate=datetime.strptime(bdate,"%Y-%m-%d")

        gender=Gender_Mune.get()
        if Address_Entry.get() is None:
            errors+=1
        address=Address_Entry.get()
        today=date.today()
        client_id=employee.Employee.generateID()
        
        for i in range(len(self.services)):
            Selected=""
            if int(self.services[i].get()) == 1:
                Selected += str(i)
                this=self.Value[i]
                # print(self.services[i].get())
                # print(this)
                chosen_serve.append(this)
                # print(self.services[i].get())

        if len(chosen_serve)==0:
            errors+=1
            
        if errors ==0:
            client=(client_id,name,age,gender,bdate,address)
            user=self.user
            user.addNewClient(user, client,chosen_serve,today)
            messagebox.showinfo("Client Entered Successfully","New Client Data Registered Successfully")
                
            self.Page_FrontDesk.destroy()
            self.Page_Dashboard.pack()

        else: messagebox.showerror("Input Error","There is one or More Errors in Data Entered. \nPlease Make sure that the Data you entered followed Specifications")

    def logout(self):
        self.Dashboard_GUI.destroy()
        interface=LoginPage.Loginpage()
        interface.start()

    def FrontDesk(self):
        width= self.Dashboard_GUI.winfo_screenwidth()
        height=self.Dashboard_GUI.winfo_screenheight()
        self.Dashboard_GUI.geometry(f'{width}x{height}+{0}+{0}')
        self.Dashboard_GUI.resizable(True,True)
        self.Page_Dashboard.forget()
        self.Page_FrontDesk=Frame(self.Dashboard_GUI)
        self.Page_FrontDesk.pack(expand=1, fill=BOTH)

        Frame_Header=Frame(self.Page_FrontDesk,width=1360,height=50,bg='#BDFFC4',highlightbackground="black",highlightthickness=1)
        Frame_Header.pack(fill=X)
        image5 = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((40, 40)))
        IMG_HEADER_FD=Label(Frame_Header,image=image5,bg='#BDFFC4',width=40,height=40)
        IMG_HEADER_FD.image=image5
        IMG_HEADER_FD.pack(side=LEFT,padx=5)
        HEADER_TITLE=Label(Frame_Header,text="City Health Office",bg='#BDFFC4',font='Roboto 25 bold').pack(side=LEFT)

        Toggle_Button=Menubutton(Frame_Header,width=5,text="=",highlightbackground="black",highlightthickness=1,justify=RIGHT)
        Toggle_Button.pack(side=RIGHT,padx=5)
        Toggle_Button.menu=Menu(Toggle_Button)
        Toggle_Button["menu"]=Toggle_Button.menu

        Toggle_Button.menu.add_command(label="Home",font='Roboto 12',command=lambda:self.HOME_PAGE(self.Close_ID[4]))
        Toggle_Button.menu.add_command(label="Logout",font='Roboto 12',command=lambda:self.logout())

        HEADER_USERNAME=Label(Frame_Header,text=str(self.user.username),bg='#BDFFC4',font='Roboto 20 ').pack(side=RIGHT)

        #Header-------
        #Body-------
        Frame_Body=Frame(self.Page_FrontDesk)
        Frame_Body.pack(side=RIGHT,fill=BOTH,expand=True)

        Frame_LIST=Frame(Frame_Body)
        Frame_LIST.pack(side=LEFT,fill=BOTH,expand=True)
        FrontDesk_Title=Label(Frame_LIST,text="New Client",font='Roboto 75')
        FrontDesk_Title.pack(pady=30)

        Space=Frame(Frame_LIST,height=150)
        Space.pack(fill=BOTH)

        Box_Title=Label(Frame_LIST,text="Laboratory Test List",width=28,anchor=W,font='Roboto 25')
        Box_Title.pack()
        self.Box=Frame(Frame_LIST,highlightbackground="black",highlightthickness=1)
        self.Box.pack()
        self.Box_R=Frame(self.Box)
        self.Box_R.pack(side=LEFT,fill=BOTH,padx=5)
        self.Box_L=Frame(self.Box)
        self.Box_L.pack(side=RIGHT,fill=BOTH)
        #checkbox
        self.Value=[
                    "Complete Blood Count",
                    "Blood Type",
                    "Stool Exam",
                    "Urinalysis (Urine Test)",
                    "Syphilis Rapid Test",
                    "Hepatitis B (Antigen Test)",
                    "Anti-HAV Test",
                    "Drug Test",
                    "Pregnancy Test",
                    "Fasting Blood Suger Test",
                    "Blood Uric Acid Test",
                    "Blood Cholesterol Test",
                    "Blood Creatinine Test",
                    "Acid Fast Staining",
                    "X-Ray Test",
                    "Serology",
                    "Medical Certificate",
                    "Fecalysis"
                    ]
        
        self.services=[]
        for i in range (len(self.Value)):
            Test=IntVar()
            Test.set(0)
            self.services.append(Test)

        Checkbutton(self.Box_R,text="Complete Blood Count",width=25,anchor=W,variable=self.services[0],font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Blood Type",variable=self.services[1],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Stool Exam",variable=self.services[2],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Urinalysis (“Urine Test”)",variable=self.services[3],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Syphilis Rapid Test",variable=self.services[4],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Hepatitis B (“Antigen Test”)",variable=self.services[5],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Anti-HAV Test",variable=self.services[6],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Drug Test",variable=self.services[7],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_R,text="Pregnancy Test",variable=self.services[8],anchor=W,font='Roboto 12 ').pack(fill=X)

        Checkbutton(self.Box_L,text="Fasting Blood Suger Test",width=25,anchor=W,variable=self.services[9],font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Blood Uric Acid Test",variable=self.services[10],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Blood Cholesterol Test",variable=self.services[11],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Blood Creatinine Test",variable=self.services[12],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Acid Fast Staining",variable=self.services[13],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="X-Ray Test",variable=self.services[14],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Serology",variable=self.services[15],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Medical Certificate",variable=self.services[16],anchor=W,font='Roboto 12 ').pack(fill=X)
        Checkbutton(self.Box_L,text="Fecalysis",variable=self.services[17],anchor=W,font='Roboto 12 ').pack(fill=X)
        #Body-------

        #IMPUT-----------------
        Frame_Input=Frame(Frame_Body)
        Frame_Input.pack(fill=BOTH,side=RIGHT,expand=TRUE)
                
        FrameIMG=Frame(Frame_Input,width=300,height=300)
        FrameIMG.pack(pady=40)

        image = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((300, 300)))
        Img=Label(FrameIMG,image = image)
        Img.image=image
        Img.pack(fill=BOTH)

        contener_FD=Frame(Frame_Input)
        contener_FD.pack(pady=20)
        global Name_Entry
        Name_Label=Label(contener_FD,text="Name: ",font='Roboto 12',anchor=W).pack(fill=X)
        Name_Entry=Entry(contener_FD,width=59,borderwidth=3,font='Roboto 12')
        Name_Entry.pack()

        contener_FD2=Frame(contener_FD)
        contener_FD2.pack(pady=4)

        global AGE_Entry
        AGE_Label=Label(contener_FD2,text="Age: ",font='Roboto 12').pack(side=LEFT)
        AGE_Entry=Entry(contener_FD2,width=5,font='Roboto 12',borderwidth=3)
        AGE_Entry.pack(side=LEFT)

        global Year_Birth
        thisyear=datetime.today().year
        Year_number=list(range(thisyear,1900,-1))
        Year=Year_number
        Year_Birth=ttk.Combobox(contener_FD2,value=Year,font='Roboto 12',width=6,state='readonly')
        Year_Birth.set(thisyear)
        Year_Birth.current(0)

        month_num=datetime.now()
        month_num=month_num.month
        curr_month=datetime.strptime(str(month_num),"%m")
        curr_year=datetime.strptime(Year_Birth.get(),"%Y")

        cal=calendar.monthcalendar(int(curr_year.year),int(curr_month.month))
        number=[day for week in cal for day in week if day != 0]
        self.client_bmonth=month_num
        # print(month_num)
        def setMonth(event):
            month_num= datetime.strptime(event.widget.get(), '%B').month
            # print(month_num)
            curr_month=datetime.strptime(str(month_num),"%m")
            curr_year=datetime.strptime(Year_Birth.get(),"%Y")
            global cal,number
            cal=calendar.monthcalendar(int(curr_year.year),int(curr_month.month))
            number=[day for week in cal for day in week if day != 0]
            Day_Birth.config(value=number)
            self.client_bmonth=month_num

        def calculate_age(event):
            birthdate=event.widget.get()
            birthdate=datetime.strptime(birthdate,'%Y')
            birthdate=birthdate.year
            today = date.today()
            age= today.year - birthdate
            AGE_Entry.delete(0,END)
            AGE_Entry.insert(0,age)

            month_num=datetime.strptime(Month_Birth.get(), '%B').month
            curr_month=datetime.strptime(str(month_num),"%m")
            curr_year=datetime.strptime(Year_Birth.get(),"%Y")
            global cal,number
            cal=calendar.monthcalendar(int(curr_year.year),int(curr_month.month))
            number=[day for week in cal for day in week if day != 0]
            Day_Birth.config(value=number)
            self.client_bmonth=month_num

        Year_Birth.bind("<<ComboboxSelected>>",calculate_age)

        Birth_Label=Label(contener_FD2,text="Birthdate:",font="Roboto 12").pack(side=LEFT)
        global Month_Birth
        Month=['January','February','March','April','May','June','July','August','September','October','November','December']
        Month_Birth=ttk.Combobox(contener_FD2,value=Month,font='Roboto 12',width=9,state='readonly')
        Month_Birth.current(0)
        
        Month_Birth.bind("<<ComboboxSelected>>",setMonth)

        global Day_Birth
        Day=number
        Day_Birth=ttk.Combobox(contener_FD2,value=Day,font='Roboto 12',width=2,state='readonly')
        Day_Birth.current(0)
        Month_Birth.pack(side=LEFT)
        Day_Birth.pack(side=LEFT)
        Year_Birth.pack(side=LEFT)

        global Gender_Mune
        Gender_Label=Label(contener_FD2,text="Gender:",font='Roboto 12').pack(side=LEFT)
        Option=["Male","Female","Other"]
        Gender_Mune=ttk.Combobox(contener_FD2,value=Option,font='Roboto 12',width=7,state='readonly')
        Gender_Mune.current(0)
        Gender_Mune.pack(side=LEFT)

        global Address_Entry,addrs
        addrs=StringVar()
        Address_Label=Label(contener_FD,text="Address: ",font='Roboto 12',anchor=W).pack(fill=X)
        Address_Entry=Entry(contener_FD,textvariable=addrs,borderwidth=3,font='Roboto 12')
        Address_Entry.pack(fill=X)

        Submit_Input=Button(contener_FD,text="Submit",width=10,bg="green",font='Roboto 11',command=self.showCheckbox)
        Submit_Input.pack(side=RIGHT,pady=20)

#FrontDesk-END>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def Record_on_close(self):
        global PageOpen
        if messagebox.askokcancel('Close', 'Are you sure you want to close the View Page all the data will not be Save?'):
            PageOpen=1
            self.RecordPage.destroy()
    
    def Record(self,value):
        global PageOpen
        if PageOpen < 2:
            self.RecordPage=Toplevel()
            self.RecordPage.title("Record")
            self.RecordPage
            Record_width=450
            Record_height=600
            self.RecordPage.geometry(f'{Record_width}x{Record_height}+{480}+{100}')
            self.RecordPage.protocol("WM_DELETE_WINDOW", self.Record_on_close)
            self.RecordPage.resizable(False,False)


            self.RecordBody= Frame(self.RecordPage)
            self.RecordBody.pack(expand=1,fill=BOTH)

            

            global Record_List
            Record_search_LB=Label(self.RecordBody,text="SEARCH: ",font='Roboto 12 bold').place(x=10,y=100)
            Record_search_EN=Entry(self.RecordBody,font='Roboto 12',borderwidth=5)
            Record_search_EN.place(x=90,y=98,relwidth=0.6)

            RecordFrame=Frame(self.RecordBody,highlightbackground="black",highlightthickness=1)
            RecordFrame.place(x=0,y=130,relwidth=1.0,relheight=0.78)
            RecordBOX= Canvas(RecordFrame,highlightbackground="black",highlightthickness=1,bg="White")
            RecordBOX.pack(side=LEFT,fill=BOTH,expand=1)

            Recordscroll=ttk.Scrollbar(RecordFrame,orient=VERTICAL,command=RecordBOX.yview)
            Recordscroll.pack(side=RIGHT,fill=Y)

            RecordBOX.configure(yscrollcommand=Recordscroll.set)
            RecordBOX.bind('<Configure>',lambda e: RecordBOX.configure(scrollregion= RecordBOX.bbox("all")))

            Record_List=Frame(RecordBOX,highlightbackground="black",highlightthickness=2)
            RecordBOX.create_window((0,0),window=Record_List,anchor=NW)
            #print(self.Value_Laboratory)
            if value == "Laboratory":
                Record_Title=Label(self.RecordBody,text="Laboratory Test Records",font='Roboto 25 bold').place(x=10,y=0)
                Record_paragraph=Label(self.RecordBody,text="All Laboratory TEST be save here and \nrecord of the Test being done ",font='Roboto 12').place(x=80,y=43)
                
                def setClient(name):
                    res=self.user.getClient_name(name)
                    print(res)
                    age=IntVar()
                    age.set(res[2])
                    self.AGE_Entry.config(textvariable=age)
                    self.Name_Entry.set(res[1])
                    self.Gender_Mune.set(res[3])

                    client_identification=IntVar()
                    client_identification.set(res[0])
                    self.ID_ENTRY.config(textvariable=client_identification)
                    res=self.user.getclientTests(res[0])
                    count=0
                    self.Test_Table.delete(*self.Test_Table.get_children())
                    for i in res:
                        self.Test_Table.insert('','end',iid=count,text=res[count][0],values=(i[3],))
                        count+=1

                def search():
                    record=self.user.getClients_labSearch(Record_search_EN.get())

                    global Record_List
                    Record_List.destroy()

                    Record_List=Frame(RecordBOX,highlightbackground="black",highlightthickness=2)
                    RecordBOX.create_window((0,0),window=Record_List,anchor=NW)
                    for i in range(len(record)):
                        Record_Number=Frame(Record_List,width=427,height=80)
                        Record_Number.grid(row=i,column=0)

                        Record_Page=Frame(Record_Number,width=250,height=50,highlightbackground="black",highlightthickness=1)
                        Record_Page.place(x=5,y=5,relwidth=0.98,relheight=0.9)

                        Number_BOX=Frame(Record_Page,width=70,height=50,bg="green",highlightbackground="black",highlightthickness=1)
                        Number_BOX.place(x=10,y=10)
                        Client_Number=Label(Number_BOX,text=record[i][0],font=("Roboto",17,"bold"),bg="green").place(x=4,y=6)

                        Client_Name=Label(Record_Page,text="NAME: "+record[i][1],font=("Roboto",12,"bold")).place(x=85,y=10)
                        Client_Test=Label(Record_Page,text="TEST: "+record[i][6],font=("Roboto",8,"bold")).place(x=85,y=30) 
                        Button(Record_Page,text='Select', command=lambda x=record[i][1]: setClient(x)).place(x=85,y=50)
                        
                Record_search_BT=Button(self.RecordBody,text="Search",font='Roboto 10',width=6,height=0,borderwidth=5,command=search)
                Record_search_BT.place(x=365,y=94)
                
                records=self.user.getClients_lab()
                #print(value)
                for i in range(len(records)):
                    Record_Number=Frame(Record_List,width=427,height=80)
                    Record_Number.grid(row=i,column=0)

                    Record_Page=Frame(Record_Number,width=250,height=50,highlightbackground="black",highlightthickness=1)
                    Record_Page.place(x=5,y=5,relwidth=0.98,relheight=0.9)

                    Number_BOX=Frame(Record_Page,width=70,height=50,highlightbackground="black",highlightthickness=1,bg="green")
                    Number_BOX.place(x=10,y=10)
                    Client_Number=Label(Number_BOX,text=records[i][0],font=("Roboto",17,"bold"),bg="green").place(x=4,y=6)

                    Client_Name=Label(Record_Page,text="NAME: "+records[i][1],font=("Roboto",12,"bold")).place(x=85,y=10)
                    Client_Test=Label(Record_Page,text="TEST: "+records[i][6],font=("Roboto",8,"bold")).place(x=85,y=30)
                    Button(Record_Page,text='Select', command=lambda x=records[i][1]: setClient(x)).place(x=85,y=50)
                    # Take_Button=Button(Record_Page,text="Take",font=("Roboto",8),width=6,height=0,borderwidth=5)
                    # Take_Button.place(x=360,y=37)
            
            elif value == "X_RAY":
                Record_Title_Xray=Label(self.RecordBody,text="Laboratory Test Records",font='Roboto 25 bold').place(x=10,y=0)
                Record_paragraph_Xray=Label(self.RecordBody,text="All X-Ray Laboratory TEST be save here and \nrecord of the Test being done ",font='Roboto 12').place(x=80,y=43)
                
                def setValue(name):
                    global client_id, test_id
                    res=self.user.getClient_name_Xray(name)
                    print(res)
                    test_id=res[-1]
                    self.Xray_Name_Entry.set(res[1])
                    self.Xray_Birth_Entry.set_date(res[4])
                    self.age.set(res[2])
                    self.Xray_Gender_Mune.set(res[3])
                    client_id=res[0]
                    

                def search():
                    record=self.user.getClients_XraySearch(Record_search_EN.get())
                    global Record_List
                    Record_List.destroy()

                    Record_List=Frame(RecordBOX,highlightbackground="black",highlightthickness=2)
                    RecordBOX.create_window((0,0),window=Record_List,anchor=NW)
                    for i in range(len(record)):
                        Record_Number=Frame(Record_List,width=427,height=80)
                        Record_Number.grid(row=i,column=0)

                        Record_Page=Frame(Record_Number,width=250,height=50,highlightbackground="black",highlightthickness=1)
                        Record_Page.place(x=5,y=5,relwidth=0.98,relheight=0.9)

                        Number_BOX=Frame(Record_Page,width=70,height=50,bg="green",highlightbackground="black",highlightthickness=1)
                        Number_BOX.place(x=10,y=10)
                        Client_Number=Label(Number_BOX,text=record[i][0],font=("Roboto",17,"bold"),bg="green").place(x=4,y=6)

                        Client_Name=Label(Record_Page,text="NAME: "+record[i][1],font=("Roboto",12,"bold")).place(x=85,y=10)
                        Client_Test=Label(Record_Page,text="TEST: "+record[i][6],font=("Roboto",8,"bold")).place(x=85,y=30)
                        Button(Record_Page,text='Select', command=lambda x=record[i][1]: setValue(x)).place(x=85,y=50)
                        
                Record_search_BT=Button(self.RecordBody,text="Search",font='Roboto 10',width=6,height=0,borderwidth=5,command=search)
                Record_search_BT.place(x=365,y=94)

                records=self.user.getClients_Xray()
                def take(e):
                    result=self.user.getClient(e)
                    self.client_name=result[1]
                    self.client_age=result[2]
                    self.client_gender=result[3]
                    self.client_bdate=result[4]
                    self.client_address=result[5]

                for i in range(len(records)):
                    X_RAY_Record_Number=Frame(Record_List,width=427,height=80)
                    X_RAY_Record_Number.grid(row=i,column=0)

                    X_RAY_Record_Page=Frame(X_RAY_Record_Number,width=250,height=50,highlightbackground="black",highlightthickness=1)
                    X_RAY_Record_Page.place(x=5,y=5,relwidth=0.98,relheight=0.9)

                    XRAY_Number_BOX=Frame(X_RAY_Record_Page,width=70,height=50,bg="green",highlightbackground="black",highlightthickness=1)
                    XRAY_Number_BOX.place(x=10,y=10)
                    XRAY_Client_Number=Label(XRAY_Number_BOX,text=records[i][0],font=("Roboto",17,"bold"),bg="green").place(x=4,y=6)

                    XRAY_Client_Name=Label(X_RAY_Record_Page,text="NAME: "+records[i][1],font=("Roboto",12,"bold")).place(x=85,y=10)
                    XRAY_Client_Test=Label(X_RAY_Record_Page,text="TEST: XRAY TEST",font=("Roboto",8,"bold")).place(x=85,y=30)
                    Button(X_RAY_Record_Page,text='Select', command=lambda x=records[i][1]: setValue(x)).place(x=85,y=50)

                    # XRAY_Take_Button=Button(X_RAY_Record_Page,text="Take",font=("Roboto",8),width=6,height=0,borderwidth=5,command=lambda e= records[i][0]:take(e))
                    # XRAY_Take_Button.place(x=360,y=37)
            
            PageOpen += 1

        else:
            messagebox.showinfo("Error","The Window is already Open!")

#Laboratory>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def Laboratory(self):
        if self.user.return_dept()!='Laboratory Department':
            messagebox.showerror("Access Denied","Only Employees from Laboratory Department or Users with Administrative Access can Access this Page")
        else: 
            self.Dashboard_GUI.resizable(True,True)
            width= self.Dashboard_GUI.winfo_screenwidth()
            height=self.Dashboard_GUI.winfo_screenheight()
            self.Dashboard_GUI.geometry(f'{width}x{height}+{0}+{0}')
            self.Page_Dashboard.forget()
            self.Page_Laboratory=Frame(self.Dashboard_GUI,bg="green")
            self.Page_Laboratory.pack(expand=1, fill=BOTH)


            Frame_Header=Frame(self.Page_Laboratory,width=1360,height=50,bg='#BDFFC4',highlightbackground="black",highlightthickness=1)
            Frame_Header.pack(fill=X)

            image4 = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((40, 40)))
            IMG_HEADER_Lab=Label(Frame_Header,image=image4,bg='#BDFFC4',width=40,height=40)
            IMG_HEADER_Lab.image=image4
            IMG_HEADER_Lab.pack(side=LEFT)
            HEADER_TITLE=Label(Frame_Header,text="City Health Office",bg='#BDFFC4',font='Roboto 25 bold').pack(side=LEFT)

            Toggle_Button=Menubutton(Frame_Header,width=5,text="=",highlightbackground="black",highlightthickness=1,justify=RIGHT)
            Toggle_Button.pack(side=RIGHT,padx=20)
            Toggle_Button.menu=Menu(Toggle_Button)
            Toggle_Button["menu"]=Toggle_Button.menu

            Toggle_Button.menu.add_command(label="Home",font='Roboto 12',command=lambda:self.HOME_PAGE(self.Close_ID[0]))
            Toggle_Button.menu.add_command(label="Logout",font='Roboto 12',command=lambda:self.logout())

            HEADER_USERNAME=Label(Frame_Header,text=str(self.user.username),bg='#BDFFC4',font='Roboto 20 ')
            HEADER_USERNAME.pack(side=RIGHT)

            #Header-------
            #BODY >> Laboratory
            Frame_Body=Frame(self.Page_Laboratory,highlightbackground="black",highlightthickness=1)
            Frame_Body.pack(fill=X)

            Patent_Detail=Frame(Frame_Body)
            Patent_Detail.pack(side=LEFT,padx=20)

            Labo_Title=Label(Patent_Detail,text="Laboratory Test!",font='Roboto 20',anchor=W)
            Labo_Title.pack(side=TOP,fill=X)

            Detail_1=Frame(Patent_Detail,bg="gray")
            Detail_1.pack(pady=5)

            ID_LABEL=Label(Detail_1,text="ID: ",font='Roboto 12').pack(side=LEFT)
            self.ID_ENTRY=Entry(Detail_1,width=8,font='Roboto 9',borderwidth=3,state='disabled')
            self.ID_ENTRY.pack(side=LEFT)

            Name_Label=Label(Detail_1,text="Name: ",font='Roboto 12').pack(side=LEFT)
            # Name_Entry=Entry(Detail_1,width=59,borderwidth=3,font='Roboto 9')
            res=self.user.getClients_all()
            names=[x[1] for x in res]
            self.Name_Entry=ttk.Combobox(Detail_1,value=names,font='Roboto 12',width=40)
            self.Name_Entry.pack(side=LEFT)

            def Name_entry_search(event):
                value=event.widget.get()
                if value!='':
                    lst=self.user.getClients_nameEntrySearch(value)
                    data=[]
                    for item in range(len(lst)):
                        if value.lower() in lst[item][1].lower():
                            data.append(lst[item][1])
                    event.widget['values']=data

                else:
                    lst=self.user.getClients_all()
                    event.widget['values']=([x[1] for x in lst])

            def setClient(event):
                res=self.user.getClient_name(self.Name_Entry.get())
                age=IntVar()
                age.set(res[2])
            
                self.Gender_Mune.set(res[3])

            self.Name_Entry.bind("<<ComboboxSelected>>",setClient)
            self.Name_Entry.bind("<KeyRelease>",Name_entry_search)

            AGE_Label=Label(Detail_1,text="Age: ",font='Roboto 12').pack(side=LEFT)
            self.AGE_Entry=Entry(Detail_1,width=8,font='Roboto 9',borderwidth=3,state='disabled')
            self.AGE_Entry.pack(side=LEFT)

            Detail_2=Frame(Patent_Detail)
            Detail_2.pack(side=BOTTOM,fill=X)

            Gender_Label=Label(Detail_2,text="Gender:",font='Roboto 12').pack(side=LEFT)
            Option=["Male","Female","Other"]
            self.Gender_Mune=ttk.Combobox(Detail_2,value=Option,font='Roboto 12',state='readonly')
            self.Gender_Mune.set("Select Gender")
            # Gender_Mune.bind("<<ComboboxSelected>>",Gender_Click)
            self.Gender_Mune.pack(side=LEFT,padx=3)
            Date_Label=Label(Detail_2,text="Date:",font=("Roboto 12")).pack(side=LEFT)
            Date_Entry=DateEntry(Detail_2,width=10,backgroud="magenta3",foreground="White",font="Roboto 12",bd=2,archor=W)
            Date_Entry.pack(side=LEFT)

            Record_Button=Button(Frame_Body,text="Record",bg="green",width=15,height=1,font=("Roboto",10),borderwidth=5,command=lambda:self.Record(self.Value_Laboratory[0]))
            Record_Button.pack(side=RIGHT)

            Testlist=Frame(Frame_Body,bg="blue",highlightbackground="black",highlightthickness=1)
            Testlist.pack(side=RIGHT)

            self.Test_Table=ttk.Treeview(Testlist,height=5)
            style=ttk.Style()
            style.theme_use("default")
            style.configure("Treeview")
            self.Test_Table['column']=("Laboratory")

            self.Test_Table.column("#0",width=0,stretch=NO)
            self.Test_Table.column("Laboratory",width=200)

            self.Test_Table.heading("#0")
            self.Test_Table.heading("Laboratory",text="Laboratory Test")
            self.Test_Table.pack(fill=Y)

            def setClient(event):
                if self.Test_Table.get_children()!=0:
                    self.Test_Table.delete(*self.Test_Table.get_children())

                res=self.user.getClient_name(self.Name_Entry.get())
                age=IntVar()
                age.set(res[2])
                self.AGE_Entry.config(textvariable=age)
                self.Gender_Mune.set(res[3])

                client_identification=IntVar()
                client_identification.set(res[0])
                self.ID_ENTRY.config(textvariable=client_identification)
                res=self.user.getclientTests(res[0])
                count=0
                for i in res:
                    self.Test_Table.insert('','end',iid=count,text=res[count][0],values=(i[3],))
                    count+=1
                
            # def double_click(event):
            #     iid=Test_Table.focus()
            #     Test_Table.set(iid, 'Complete','Done')
            #     item=Test_Table.item(iid)["text"]
            #     self.user.markTest_as_done(item)


            self.Name_Entry.bind("<<ComboboxSelected>>",setClient)
            #Frame for the Testing 
            Frame_Test=Frame(self.Page_Laboratory,highlightbackground="black",highlightthickness=1,bg="blue")
            Frame_Test.pack(expand=1,fill=BOTH)

            Header_Test=Frame(Frame_Test,highlightbackground="black",highlightthickness=1)
            Header_Test.pack(fill=X)

            Test_Label=Label(Header_Test,text="Laboratory Test",font="Roboto 15",anchor=W)
            Test_Label.pack(side=LEFT)

            Test=[  
                    "Complete Blood Count / Hematology",
                    "Urinalysis",
                    "Serology",
                    "Miscelaneous",
                    "Fecalysis"
                ]

            def Option_TEST(event):
                if LabTest_Mune.get() == "Serology":
                    Miscelaneous_Page.pack_forget()
                    Urinalysis_Page.pack_forget()
                    CBC_Page.pack_forget()
                    FE_Page.pack_forget()

                    Serology_Page.pack(expand=1,fill=BOTH)

                    Serology_Title = Label(Serology_Page,text=LabTest_Mune.get(),font=("Roboto",20,"bold"))
                    Serology_Title.place(x=570,y=30)

                    #RIGHT
                    ST_Box=Frame(Serology_Page,bg='white')
                    ST_Box.place(x=380,y=150)
                    ST_BOX1_R= Label(ST_Box,text="",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX1_R.grid(row=0,column=0)
                    ST_BOX2_R= Label(ST_Box,text="BLOOD TYPE",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX2_R.grid(row=1,column=0)
                    ST_BOX3_R= Label(ST_Box,text="HEPATITIS B SCREENING (HBsAg)",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX3_R.grid(row=2,column=0)
                    ST_BOX4_R= Label(ST_Box,text="ANTI-HAV SCREENING (HAV lgG/igM)",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX4_R.grid(row=3,column=0)
                    ST_BOX5_R= Label(ST_Box,text="SYPHILIS SCREENING",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX5_R.grid(row=4,column=0)
                    ST_BOX6_R= Label(ST_Box,text="DENGUE NS1 ANTIGEN TEST",font=("Roboto",15,"bold"),width=35,anchor=W,highlightbackground="black",highlightthickness=1)
                    ST_BOX6_R.grid(row=5,column=0)
                    
                    #LEFT
                    ST_BOX7_L= Label(ST_Box,text="RESULT",font=("Roboto",15,"bold"),width=18,highlightbackground="black",highlightthickness=1)
                    ST_BOX7_L.grid(row=0,column=1)
                    ST_BOX8_L= Entry(ST_Box,text="",font=("Roboto",15,"bold"),borderwidth=3,highlightbackground="black",highlightthickness=1)
                    ST_BOX8_L.grid(row=1,column=1)
                    ST_BOX9_L= Entry(ST_Box,text="",font=("Roboto",15,"bold"),borderwidth=3,highlightbackground="black",highlightthickness=1)
                    ST_BOX9_L.grid(row=2,column=1)
                    ST_BOX10_L= Entry(ST_Box,text="",font=("Roboto",15,"bold"),borderwidth=3,highlightbackground="black",highlightthickness=1)
                    ST_BOX10_L.grid(row=3,column=1)
                    ST_BOX11_L= Entry(ST_Box,text="",font=("Roboto",15,"bold"),borderwidth=3,highlightbackground="black",highlightthickness=1)
                    ST_BOX11_L.grid(row=4,column=1)
                    ST_BOX12_L= Entry(ST_Box,text="",font=("Roboto",15,"bold"),borderwidth=3,highlightbackground="black",highlightthickness=1)
                    ST_BOX12_L.grid(row=5,column=1)

                    def submit():
                        serviceid=self.user.get_test_id("Serology")
                        client_id=self.user.getClient_name(Name_Entry.get())
                        services=self.user.getClientTestRequests(client_id,serviceid)
                        if tuple(serviceid) not in services:
                            messagebox.showerror("Error","This Test was not Requested by the Client")
                        else: 
                            blood_type=ST_BOX8_L.get()
                            hepatitis_b_Screening=ST_BOX9_L.get()
                            anti_hav_screening=ST_BOX10_L.get()
                            syphilis_screen=ST_BOX11_L.get()
                            dengue_ns1_antigen_test=ST_BOX12_L.get()
                            document=Path(__file__).parent / os.getenv("SEROLOGY_TEMPLATE")
                            doc=DocxTemplate(document)
                            OR_Num=self.user.generateClient_ORNumber()
                            context={
                                "NAME":Name_Entry.get(),
                                "AGE_SEX":AGE_Entry.get()+'/'+Gender_Mune.get(),
                                "DATE":self.test_date,
                                "OR_NO":OR_Num,
                                "BLOODTYPE": blood_type,
                                "HEPA_B_SCREEN": hepatitis_b_Screening,
                                "ANTI_HAV_SCREEN": anti_hav_screening,
                                "SYPHILIS_SCREEN": syphilis_screen,
                                "DENGUE_ANTIGEN_TEST": dengue_ns1_antigen_test,
                                "MEDTECH_NAME":self.user.fname+" "+self.user.lname,
                                "PATHOLOGIST":os.getenv("PATHOLOGIST")
                            }
                            doc.render(context)
                            path=os.getenv('DRIVE_PATH')+"/Serology"
                            Exists=os.path.exists(path)
                            if not Exists:
                                os.makedirs(path)

                            doc.save(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx")

                            win32api.ShellExecute(0, "print", path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx", None, ".", 0)
                            serviceid=self.user.get_test_id("Serology")
                            total=self.user.get_test_price(serviceid[0])
                            id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                            self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                            test_id=self.user.get_tests_id(client_id[0],serviceid[0])
                            self.user.markTest_as_done(test_id[0][0])


                    ST_Button=Button(Serology_Page,text="Submit",font=("Roboto",10,"bold"),width=10,height=1,borderwidth=5,command=lambda: submit())
                    ST_Button.place(x=1200,y=430)

                elif LabTest_Mune.get() == "Miscelaneous":
                    Serology_Page.pack_forget()
                    Urinalysis_Page.pack_forget()
                    CBC_Page.pack_forget()
                    FE_Page.pack_forget()

                    Miscelaneous_Page.pack(expand=1,fill=BOTH)
                    Miscelaneous_Title = Label(Miscelaneous_Page,text=LabTest_Mune.get(),font=("Roboto",20,"bold"))
                    Miscelaneous_Title.place(x=570,y=30)

                    PT_Box=Frame(Miscelaneous_Page,bg='white')
                    PT_Box.place(x=430,y=200)
                    PT_BOX1= Label(PT_Box,text="TEST",width=20,anchor=W,font=("Roboto",15,"bold"),highlightbackground="black",highlightthickness=1)
                    PT_BOX1.grid(row=0,column=0)
                    # PT_BOX2= Label(PT_Box,text="PREGNANCY TEST",width=20,anchor=W,font=("Roboto",15,"bold"))
                    res_test=self.user.getAllTest()
                    TEST=[x[0] for x in res_test]
                    PT_BOX2= ttk.Combobox(PT_Box,value=TEST,font=("Roboto",15),state='readonly')
                    PT_BOX2.grid(row=0,column=1)
                    PT_BOX3= Label(PT_Box,text="RESULT",width=20,anchor=W,font=("Roboto",15,"bold"),highlightbackground="black",highlightthickness=1)
                    PT_BOX3.grid(row=1,column=0)

                    # PT_Result=["POSITIVE","NEGATIVE"]
                    # PT_BOX4=ttk.Combobox(PT_Box,value=PT_Result,font=("Roboto",15),state='readonly')
                    # PT_BOX4.set("Select Result")
                    PT_BOX4=Entry(PT_Box,font=("Roboto",12),width=25,borderwidth=3,highlightbackground="black",highlightthickness=1)
                    PT_BOX4.grid(row=1,column=1)

                    def submit():
                        print(PT_BOX2.get())
                        print(res[0][0])
                        serviceid=self.user.get_test_id(PT_BOX2.get())
                        services=self.user.getClientTestRequests(int(self.ID_ENTRY.get()),serviceid)
                        print(serviceid)
                        print(services)
                        if tuple(serviceid) not in services:
                            messagebox.showerror("Error","This Test was not Requested by the Client")
                        else: 
                            document=Path(__file__).parent / os.getenv("MISCELLANEOUS_TEMPLATE")
                            doc=DocxTemplate(document)
                            OR_Num=self.user.generateClient_ORNumber()
                            context={
                                "NAME":Name_Entry.get(),
                                "AGE_SEX":AGE_Entry.get()+'/'+Gender_Mune.get(),
                                "DATE":self.test_date,
                                "OR_NO":OR_Num,
                                "TEST":PT_BOX2.get(),
                                "RESULT":PT_BOX4.get(),
                                "MEDTECH_NAME":self.user.fname+" "+self.user.lname,
                                "PATHOLOGIST":os.getenv("PATHOLOGIST")
                            }
                            doc.render(context)
                            path=os.getenv('DRIVE_PATH')+"/"+PT_BOX2.get()
                            Exists=os.path.exists(path)
                            if not Exists:
                                os.makedirs(path)

                            doc.save(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx")
                            win32api.ShellExecute(0, "print", path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx", None, ".", 0)
                            total=self.user.get_test_price(serviceid[0])
                            id=self.user.save_to_summary(total[0],serviceid[0],int(self.ID_ENTRY.get()))
                            self.user.update_summaryID_test(id,int(self.ID_ENTRY.get()),serviceid[0])
                            test_id=self.user.get_tests_id(int(self.ID_ENTRY.get()),serviceid[0])
                            self.user.markTest_as_done(test_id[0][0])

                    PT_Button=Button(Miscelaneous_Page,text="Submit",font=("Roboto",10,"bold"),width=10,height=1,borderwidth=5, command=lambda: submit())
                    PT_Button.place(x=1200,y=430)
                
                elif LabTest_Mune.get() == "Urinalysis":
                    Serology_Page.pack_forget()
                    Miscelaneous_Page.pack_forget()
                    CBC_Page.pack_forget()
                    FE_Page.pack_forget()

                    Urinalysis_Page.pack(expand=1,fill=BOTH)

                    Urinalysis_Title = Label(Urinalysis_Page,text=LabTest_Mune.get(),font=("Roboto",20,"bold"))
                    Urinalysis_Title.place(x=570,y=30)

                    UR_Box=Frame(Urinalysis_Page,bg='white')
                    UR_Box.place(x=130,y=90)
                    UR_Column1_BOX= Label(UR_Box,text="COLOR",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column1_BOX.grid(row=0,column=0)
                    UR_Column1_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column1_BOX1.grid(row=0,column=1,padx=1)
                    UR_Column1_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column1_BOX2.grid(row=0,column=2)

                    UR_Column2_BOX= Label(UR_Box,text="CLARITY",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column2_BOX.grid(row=1,column=0)
                    UR_Column2_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column2_BOX1.grid(row=1,column=1,padx=1)
                    UR_Column2_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column2_BOX2.grid(row=1,column=2)

                    UR_Column3_BOX= Label(UR_Box,text="BLOOD",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column3_BOX.grid(row=2,column=0)
                    UR_Column3_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column3_BOX1.grid(row=2,column=1,padx=1)
                    UR_Column3_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column3_BOX2.grid(row=2,column=2)

                    UR_Column4_BOX= Label(UR_Box,text="BILIRUBIN",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column4_BOX.grid(row=4,column=0)
                    UR_Column4_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column4_BOX1.grid(row=4,column=1,padx=1)
                    UR_Column4_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column4_BOX2.grid(row=4,column=2)

                    UR_Column5_BOX= Label(UR_Box,text="LEUKOCYTE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column5_BOX.grid(row=5,column=0)
                    UR_Column5_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column5_BOX1.grid(row=5,column=1,padx=1)
                    UR_Column5_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column5_BOX2.grid(row=5,column=2)

                    UR_Column6_BOX= Label(UR_Box,text="KETONE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column6_BOX.grid(row=6,column=0)
                    UR_Column6_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column6_BOX1.grid(row=6,column=1,padx=1)
                    UR_Column6_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column6_BOX2.grid(row=6,column=2)

                    UR_Column7_BOX= Label(UR_Box,text="NITRITE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column7_BOX.grid(row=7,column=0)
                    UR_Column7_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column7_BOX1.grid(row=7,column=1,padx=1)
                    UR_Column7_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column7_BOX2.grid(row=7,column=2)

                    UR_Column8_BOX= Label(UR_Box,text="PROTEIN",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column8_BOX.grid(row=8,column=0)
                    UR_Column8_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column8_BOX1.grid(row=8,column=1,padx=1)
                    UR_Column8_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column8_BOX2.grid(row=8,column=2)

                    UR_Column9_BOX= Label(UR_Box,text="GLUCOSE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column9_BOX.grid(row=9,column=0)
                    UR_Column9_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column9_BOX1.grid(row=9,column=1,padx=1)
                    UR_Column9_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column9_BOX2.grid(row=9,column=2)

                    UR_Column10_BOX= Label(UR_Box,text="PH",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column10_BOX.grid(row=10,column=0)
                    UR_Column10_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column10_BOX1.grid(row=10,column=1,padx=1)
                    UR_Column10_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column10_BOX2.grid(row=10,column=2)

                    UR_Column11_BOX= Label(UR_Box,text="SPECIFIC GRAVITY",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column11_BOX.grid(row=11,column=0)
                    UR_Column11_BOX1= Entry(UR_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column11_BOX1.grid(row=11,column=1,padx=1)
                    UR_Column11_BOX2= Label(UR_Box,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column11_BOX2.grid(row=11,column=2)

                    UR_Box_Side=Frame(Urinalysis_Page,bg='white')
                    UR_Box_Side.place(x=650,y=90)

                    UR_Column12_BOX= Label(UR_Box_Side,text="MICROSCOPIC EXAMINATION",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column12_BOX.grid(row=12,column=0)
                    UR_Column12_BOX1= Label(UR_Box_Side,text="RESULT",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column12_BOX1.grid(row=12,column=1,padx=1)
                    UR_Column12_BOX2= Label(UR_Box_Side,text="UNIT",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column12_BOX2.grid(row=12,column=2)

                    UR_Column13_BOX= Label(UR_Box_Side,text="WBC",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column13_BOX.grid(row=13,column=0)
                    UR_Column13_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column13_BOX1.grid(row=13,column=1,padx=1)
                    UR_Column13_BOX2= Label(UR_Box_Side,text="/HPF",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column13_BOX2.grid(row=13,column=2)

                    UR_Column14_BOX= Label(UR_Box_Side,text="RBC",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column14_BOX.grid(row=14,column=0)
                    UR_Column14_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column14_BOX1.grid(row=14,column=1,padx=1)
                    UR_Column14_BOX2= Label(UR_Box_Side,text="/HPF",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column14_BOX2.grid(row=14,column=2)

                    UR_Column15_BOX= Label(UR_Box_Side,text="EPITHELIAL CELLS",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column15_BOX.grid(row=15,column=0)
                    UR_Column15_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column15_BOX1.grid(row=15,column=1,padx=1)
                    UR_Column15_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column15_BOX2.grid(row=15,column=2)

                    UR_Column16_BOX= Label(UR_Box_Side,text="MUCOUS THREADS",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column16_BOX.grid(row=16,column=0)
                    UR_Column16_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column16_BOX1.grid(row=16,column=1,padx=1)
                    UR_Column16_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column16_BOX2.grid(row=16,column=2)

                    UR_Column17_BOX= Label(UR_Box_Side,text="BACTERIA",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column17_BOX.grid(row=17,column=0)
                    UR_Column17_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column17_BOX1.grid(row=17,column=1,padx=1)
                    UR_Column17_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column17_BOX2.grid(row=17,column=2)

                    UR_Column18_BOX= Label(UR_Box_Side,text="A. URATES / PHOSPHATE",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column18_BOX.grid(row=18,column=0)
                    UR_Column18_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column18_BOX1.grid(row=18,column=1,padx=1)
                    UR_Column18_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column18_BOX2.grid(row=18,column=2)

                    UR_Column19_BOX= Label(UR_Box_Side,text="CASTS",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column19_BOX.grid(row=19,column=0)
                    UR_Column19_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column19_BOX1.grid(row=19,column=1,padx=1)
                    UR_Column19_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column19_BOX2.grid(row=19,column=2)

                    UR_Column20_BOX= Label(UR_Box_Side,text="CRYSTALS:",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column20_BOX.grid(row=20,column=0)
                    UR_Column20_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column20_BOX1.grid(row=20,column=1,padx=1)
                    UR_Column20_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column20_BOX2.grid(row=20,column=2)

                    UR_Column21_BOX= Label(UR_Box_Side,text="OTHERS",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column21_BOX.grid(row=21,column=0)
                    UR_Column21_BOX1= Entry(UR_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    UR_Column21_BOX1.grid(row=21,column=1,padx=1)
                    UR_Column21_BOX2= Label(UR_Box_Side,text=" ",width=18,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    UR_Column21_BOX2.grid(row=21,column=2)

                    def submit():
                        client_id=self.user.getClient_name(Name_Entry.get())
                        serviceid=self.user.get_test_id("Urinalysis (Urine Test)")
                        services=self.user.getClientTestRequests(client_id,serviceid)
                        if tuple(serviceid) not in services:
                            messagebox.showerror("Error","This Test was not Requested by the Client")
                        else: 
                            document=Path(__file__).parent / os.getenv("URINALYSIS_TEMPLATE")
                            doc=DocxTemplate(document)
                            OR_Num=self.user.generateClient_ORNumber()
                            context={
                                "NAME":Name_Entry.get(),
                                "AGE_SEX":AGE_Entry.get()+'/'+Gender_Mune.get(),
                                "DATE":self.test_date,
                                "OR_NO":OR_Num,

                                "COLOR":UR_Column1_BOX1.get(),
                                "CLARITY":UR_Column2_BOX1.get(),
                                "BLOOD":UR_Column3_BOX1.get(),
                                "BILIRUBIN":UR_Column4_BOX1.get(),
                                "LEUKOCYTE":UR_Column5_BOX1.get(),
                                "KETONE":UR_Column6_BOX1.get(),
                                "NITRITE":UR_Column7_BOX1.get(),
                                "PROTEIN":UR_Column8_BOX1.get(),
                                "GLUCOSE":UR_Column9_BOX1.get(),
                                "PH":UR_Column10_BOX1.get(),
                                "SPECIFIC_GRAVITY":UR_Column11_BOX1.get(),
                                "WBC":UR_Column13_BOX1.get(),
                                "RBC":UR_Column14_BOX1.get(),
                                "EPITHERIAL_CELLS":UR_Column15_BOX1.get(),
                                "MUCOUS_THREADS":UR_Column16_BOX1.get(),
                                "BACTERIA":UR_Column17_BOX1.get(),
                                "PHOSPHATE":UR_Column18_BOX1.get(),
                                "CASTS":UR_Column19_BOX1.get(),
                                "CRYSTALS":UR_Column20_BOX1.get(),
                                "OTHERS":UR_Column21_BOX1.get(),
                                
                                "MEDTECH_NAME":self.user.fname+" "+self.user.lname,
                                "PATHOLOGIST":os.getenv("PATHOLOGIST")
                            }
                            doc.render(context)
                            path=os.getenv('DRIVE_PATH')+"/Urinalysis"
                            Exists=os.path.exists(path)
                            if not Exists:
                                os.makedirs(path)

                            doc.save(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx")
                            win32api.ShellExecute(0, "print", path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx", None, ".", 0)
                            total=self.user.get_test_price(serviceid[0])
                            id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                            self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                            test_id=self.user.get_tests_id(client_id[0],serviceid[0])
                            self.user.markTest_as_done(test_id[0][0])

                    ST_Button=Button(Urinalysis_Page,text="Submit",font=("Roboto",10,"bold"),width=10,height=1,borderwidth=5,command=submit)
                    ST_Button.place(x=1200,y=430)
                
                elif LabTest_Mune.get() == "Complete Blood Count / Hematology":
                    Serology_Page.pack_forget()
                    Miscelaneous_Page.pack_forget()
                    Urinalysis_Page.pack_forget()
                    FE_Page.pack_forget()

                    CBC_Page.pack(expand=1,fill=BOTH)
                    CBC_Title = Label(CBC_Page,text=LabTest_Mune.get(),font=("Roboto",20,"bold"))
                    CBC_Title.place(x=400,y=30)

                    CBC_Box=Frame(CBC_Page,bg='white')
                    CBC_Box.place(x=60,y=120)
                    CBC_Column1_BOX= Label(CBC_Box,text="",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column1_BOX.grid(row=0,column=0)
                    CBC_Column1_BOX1= Label(CBC_Box,text="RESULT",width=18,anchor=CENTER,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column1_BOX1.grid(row=0,column=1,padx=1)
                    CBC_Column1_BOX2= Label(CBC_Box,text="NORMAL VALUES",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column1_BOX2.grid(row=0,column=2)

                    CBC_Column2_BOX= Label(CBC_Box,text="WBC",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column2_BOX.grid(row=1,column=0)
                    CBC_Column2_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column2_BOX1.grid(row=1,column=1,padx=1)
                    CBC_Column2_BOX2= Label(CBC_Box,text="5,000 - 10,000 / CUMM",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column2_BOX2.grid(row=1,column=2)

                    CBC_Column3_BOX= Label(CBC_Box,text="RBC",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column3_BOX.grid(row=2,column=0)
                    CBC_Column3_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column3_BOX1.grid(row=2,column=1,padx=1)
                    CBC_Column3_BOX2= Label(CBC_Box,text="F: 3.8 - 5.1 10^ uL ; M: 4.20 - 5.6 10^ uL",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column3_BOX2.grid(row=2,column=2)

                    CBC_Column4_BOX= Label(CBC_Box,text="HEMOGLOBIN",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column4_BOX.grid(row=4,column=0)
                    CBC_Column4_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column4_BOX1.grid(row=4,column=1,padx=1)
                    CBC_Column4_BOX2= Label(CBC_Box,text="F: 11.70 - 14.5g/dL ; M: 13.7 - 16.7g/dL",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column4_BOX2.grid(row=4,column=2)

                    CBC_Column5_BOX= Label(CBC_Box,text="HEMATOCRIT",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column5_BOX.grid(row=5,column=0)
                    CBC_Column5_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column5_BOX1.grid(row=5,column=1,padx=1)
                    CBC_Column5_BOX2= Label(CBC_Box,text="F: 34.1  - 44.3vol% ; M: 13.7 - 49.7 vol%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column5_BOX2.grid(row=5,column=2)

                    CBC_Column6_BOX= Label(CBC_Box,text="MCV",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column6_BOX.grid(row=6,column=0)
                    CBC_Column6_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column6_BOX1.grid(row=6,column=1,padx=1)
                    CBC_Column6_BOX2= Label(CBC_Box,text="80-100 El",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column6_BOX2.grid(row=6,column=2)

                    CBC_Column7_BOX= Label(CBC_Box,text="MCH",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column7_BOX.grid(row=7,column=0)
                    CBC_Column7_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column7_BOX1.grid(row=7,column=1,padx=1)
                    CBC_Column7_BOX2= Label(CBC_Box,text="29 + 2 pg",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column7_BOX2.grid(row=7,column=2)

                    CBC_Column8_BOX= Label(CBC_Box,text="MCHC",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column8_BOX.grid(row=8,column=0)
                    CBC_Column8_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column8_BOX1.grid(row=8,column=1,padx=1)
                    CBC_Column8_BOX2= Label(CBC_Box,text="33.4-35.5 g/dL",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column8_BOX2.grid(row=8,column=2)

                    CBC_Column9_BOX= Label(CBC_Box,text="RDW",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column9_BOX.grid(row=9,column=0)
                    CBC_Column9_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column9_BOX1.grid(row=9,column=1,padx=1)
                    CBC_Column9_BOX2= Label(CBC_Box,text="12% to 15%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column9_BOX2.grid(row=9,column=2)

                    CBC_Column10_BOX= Label(CBC_Box,text="PLATELET",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column10_BOX.grid(row=10,column=0)
                    CBC_Column10_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column10_BOX1.grid(row=10,column=1,padx=1)
                    CBC_Column10_BOX2= Label(CBC_Box,text="150,000 - 450,000 uL",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column10_BOX2.grid(row=10,column=2)

                    CBC_Column11_BOX= Label(CBC_Box,text="MPV",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column11_BOX.grid(row=11,column=0)
                    CBC_Column11_BOX1= Entry(CBC_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBC_Column11_BOX1.grid(row=11,column=1,padx=1)
                    CBC_Column11_BOX2= Label(CBC_Box,text="8.9 - 11.8 fL",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBC_Column11_BOX2.grid(row=11,column=2)


                    #SIDE 
                    CBC_Box_Side=Frame(CBC_Page,bg='white')
                    CBC_Box_Side.place(x=700,y=120)
                    CBCS_Column1_BOX= Label(CBC_Box_Side,text="DIFFERENTIAL COUNT",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column1_BOX.grid(row=0,column=0)
                    CBCS_Column1_BOX1= Label(CBC_Box_Side,text="",width=18,anchor=CENTER,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column1_BOX1.grid(row=0,column=1,padx=1)
                    CBCS_Column1_BOX2= Label(CBC_Box_Side,text="",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column1_BOX2.grid(row=0,column=2)

                    CBCS_Column2_BOX= Label(CBC_Box_Side,text="NEUTROPHIL",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column2_BOX.grid(row=1,column=0)
                    CBCS_Column2_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column2_BOX1.grid(row=1,column=1,padx=1)
                    CBCS_Column2_BOX2= Label(CBC_Box_Side,text="45%"+" - "+"70%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column2_BOX2.grid(row=1,column=2)

                    CBCS_Column3_BOX= Label(CBC_Box_Side,text="LYMPHOCYTE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column3_BOX.grid(row=2,column=0)
                    CBCS_Column3_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column3_BOX1.grid(row=2,column=1,padx=1)
                    CBCS_Column3_BOX2= Label(CBC_Box_Side,text="18%"+" - ""45%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column3_BOX2.grid(row=2,column=2)

                    CBCS_Column4_BOX= Label(CBC_Box_Side,text="MONOCYTE",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column4_BOX.grid(row=4,column=0)
                    CBCS_Column4_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column4_BOX1.grid(row=4,column=1,padx=1)
                    CBCS_Column4_BOX2= Label(CBC_Box_Side,text="4% "+" - "+"8%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column4_BOX2.grid(row=4,column=2)

                    CBCS_Column5_BOX= Label(CBC_Box_Side,text="EOSINOPHIL",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column5_BOX.grid(row=5,column=0)
                    CBCS_Column5_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column5_BOX1.grid(row=5,column=1,padx=1)
                    CBCS_Column5_BOX2= Label(CBC_Box_Side,text="2% "+" - "+"3%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column5_BOX2.grid(row=5,column=2)

                    CBCS_Column6_BOX= Label(CBC_Box_Side,text="BASOPHIL",width=20,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column6_BOX.grid(row=6,column=0)
                    CBCS_Column6_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column6_BOX1.grid(row=6,column=1,padx=1)
                    CBCS_Column6_BOX2= Label(CBC_Box_Side,text="0% "+" - "+"2%",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column6_BOX2.grid(row=6,column=2)

                    CBCS_Column7_BOX= Label(CBC_Box_Side,text="TOTAL:",width=20,anchor=E,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column7_BOX.grid(row=7,column=0)
                    CBCS_Column7_BOX1= Entry(CBC_Box_Side,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    CBCS_Column7_BOX1.grid(row=7,column=1,padx=1)
                    CBCS_Column7_BOX2= Label(CBC_Box_Side,text="",width=33,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    CBCS_Column7_BOX2.grid(row=7,column=2)

                    def submit():
                        serviceid=self.user.get_test_id("Complete Blood Count")
                        client_id=self.user.getClient_name(Name_Entry.get())
                        services=self.user.getClientTestRequests(client_id,serviceid)
                        if tuple(serviceid) not in services:
                            messagebox.showerror("Error","This Test was not Requested by the Client")
                        else: 
                            document=Path(__file__).parent / os.getenv("CBC_TEMPLATE")
                            doc=DocxTemplate(document)
                            OR_Num=self.user.generateClient_ORNumber()
                            context={
                                "NAME":Name_Entry.get(),
                                "AGE_SEX":AGE_Entry.get()+'/'+Gender_Mune.get(),
                                "DATE":self.test_date,
                                "OR_NUM":OR_Num,
                                "WBC":CBC_Column2_BOX1.get(),
                                "RBC":CBC_Column3_BOX1.get(),
                                "HEMOGLOBIN":CBC_Column4_BOX1.get(),
                                "HEMATOCRIT":CBC_Column5_BOX1.get(),
                                "MCV":CBC_Column6_BOX1.get(),
                                "MCH":CBC_Column7_BOX1.get(),
                                "MCHC":CBC_Column8_BOX1.get(),
                                "RDW":CBC_Column9_BOX1.get(),
                                "PLATELET":CBC_Column10_BOX1.get(),
                                "MPV":CBC_Column11_BOX1.get(),
                                "NEUTROPHIL":CBCS_Column2_BOX1.get(),
                                "LYMPHOCYTE":CBCS_Column3_BOX1.get(),
                                "MONOCYTE":CBCS_Column4_BOX1.get(),
                                "EOSINOPHIL":CBCS_Column5_BOX1.get(),
                                "BASOPHIL":CBCS_Column6_BOX1.get(),
                                "TOTAL":CBCS_Column7_BOX1.get(),
                                "MEDTECH_NAME":self.user.fname+" "+self.user.lname,
                                "LICENSE_NO":self.user.license_no,
                                "PATHOLOGIST":os.getenv("PATHOLOGIST")
                            }
                            doc.render(context)
                            path=os.getenv('DRIVE_PATH')+"/Complete Blood Count"
                            Exists=os.path.exists(path)
                            if not Exists:
                                os.makedirs(path)

                            doc.save(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx")
                            win32api.ShellExecute(0, "print", path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx", None, ".", 0)
                            total=self.user.get_test_price(serviceid[0])
                            id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                            self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                            test_id=self.user.get_tests_id(client_id[0],serviceid[0])
                            self.user.markTest_as_done(test_id[0][0])

                    CBC_Button=Button(CBC_Page,text="Submit",font=("Roboto",10,"bold"),width=10,height=1,borderwidth=5,command=submit)
                    CBC_Button.place(x=1200,y=430)
                
                elif LabTest_Mune.get() == "Fecalysis":
                    Serology_Page.pack_forget()
                    CBC_Page.pack_forget()
                    Miscelaneous_Page.pack_forget()
                    Urinalysis_Page.pack_forget() 

                    FE_Page.pack(expand=1,fill=BOTH)
                    FE_Title = Label(FE_Page,text=LabTest_Mune.get(),font=("Roboto",20,"bold"))
                    FE_Title.place(x=550,y=30)

                    FE_Box=Frame(FE_Page,bg='white')
                    FE_Box.place(x=450,y=100)
                    FE_Column1_BOX= Label(FE_Box,text="COLOR",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column1_BOX.grid(row=0,column=0)
                    FE_Column1_BOX1=Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column1_BOX1.grid(row=0,column=1,padx=1)

                    FE_Column2_BOX= Label(FE_Box,text="CONSISTENCY",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column2_BOX.grid(row=1,column=0)
                    FE_Column2_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column2_BOX1.grid(row=1,column=1,padx=1)

                    FE_Column3_BOX= Label(FE_Box,text="MICROSCOPIC EXAMINATION",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column3_BOX.grid(row=2,column=0)
                    FE_Column3_BOX1= Label(FE_Box,text="RESULT",width=18,anchor=CENTER,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column3_BOX1.grid(row=2,column=1,padx=1)

                    FE_Column4_BOX= Label(FE_Box,text="WBC",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column4_BOX.grid(row=4,column=0)
                    FE_Column4_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column4_BOX1.grid(row=4,column=1,padx=1)

                    FE_Column5_BOX= Label(FE_Box,text="RBC",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column5_BOX.grid(row=5,column=0)
                    FE_Column5_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column5_BOX1.grid(row=5,column=1,padx=1)

                    FE_Column6_BOX= Label(FE_Box,text="BACTERIA",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column6_BOX.grid(row=6,column=0)
                    FE_Column6_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column6_BOX1.grid(row=6,column=1,padx=1)

                    FE_Column7_BOX= Label(FE_Box,text="FAT GLOBULES",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column7_BOX.grid(row=7,column=0)
                    FE_Column7_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column7_BOX1.grid(row=7,column=1,padx=1)

                    FE_Column8_BOX= Label(FE_Box,text="OVA OR PARASITE",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column8_BOX.grid(row=8,column=0)
                    FE_Column8_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column8_BOX1.grid(row=8,column=1,padx=1)

                    FE_Column9_BOX= Label(FE_Box,text="E. Histolytica CYST",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column9_BOX.grid(row=9,column=0)
                    FE_Column9_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column9_BOX1.grid(row=9,column=1,padx=1)

                    FE_Column10_BOX= Label(FE_Box,text="E. coli CYST",width=25,anchor=W,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1)
                    FE_Column10_BOX.grid(row=10,column=0)
                    FE_Column10_BOX1= Entry(FE_Box,width=20,font=("Roboto",10,"bold"),highlightbackground="black",highlightthickness=1,borderwidth=3)
                    FE_Column10_BOX1.grid(row=10,column=1,padx=1)

                    def submit():
                        serviceid=self.user.get_test_id("Fecalysis")
                        client_id=self.user.getClient_name(Name_Entry.get())
                        services=self.user.getClientTestRequests(client_id,serviceid)
                        if tuple(serviceid) not in services:
                            messagebox.showerror("Error","This Test was not Requested by the Client")
                        else: 
                            document=Path(__file__).parent / os.getenv("FECALYSIS_TEMPLATE")
                            doc=DocxTemplate(document)
                            OR_Num=self.user.generateClient_ORNumber()    
                            context={
                                "NAME":Name_Entry.get(),
                                "AGE_SEX":AGE_Entry.get()+'/'+Gender_Mune.get(),
                                "DATE":self.test_date,
                                "OR_NO":OR_Num,

                                "COLOR":FE_Column1_BOX1.get(),
                                "CONSISTENCY":FE_Column2_BOX1.get(),
                                "WBC": FE_Column4_BOX1.get(),
                                "RBC":FE_Column5_BOX1.get(),
                                "BACTERIA":FE_Column6_BOX1.get(),
                                "FAT_GLOBULES":FE_Column7_BOX1.get(),
                                "OVA_PARASITE":FE_Column8_BOX1.get(),
                                "E_HISTOLYTICA":FE_Column9_BOX1.get(),
                                "E_COLI":FE_Column10_BOX1.get(),

                                "MEDTECH":self.user.fname+" "+self.user.lname,
                                "PATHOLOGIST":os.getenv("PATHOLOGIST")
                            }
                            doc.render(context)
                            path=os.getenv('DRIVE_PATH')+"/Fecalysis"
                            Exists=os.path.exists(path)
                            if not Exists:
                                os.makedirs(path)

                            doc.save(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx")
                            win32api.ShellExecute(0, "print", str(path+"/"+Name_Entry.get()+"_"+str(OR_Num)+".docx"), None, ".", 0)
                            total=self.user.get_test_price(serviceid[0])
                            id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                            self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                            test_id=self.user.get_tests_id(client_id[0],serviceid[0])
                            self.user.markTest_as_done(test_id[0][0])

                    FE_Button=Button(FE_Page,text="Submit",font=("Roboto",10,"bold"),width=10,height=1,borderwidth=5, command=submit)
                    FE_Button.place(x=1200,y=430)
            
            Nothing=Label(Header_Test,text="",font='Roboto 12 bold').pack(side=RIGHT,padx=5)
            LabTest_Mune=ttk.Combobox(Header_Test,value=Test,font='Roboto 12',state='readonly')
            # LabTest_Mune.set("Serology")
            LabTest_Mune.bind("<<ComboboxSelected>>",Option_TEST)
            LabTest_Mune.pack(side=RIGHT)
            Test_Label=Label(Header_Test,text="TEST:",font='Roboto 12 bold').pack(side=RIGHT)

            Contener = Frame(Frame_Test,highlightbackground="white",highlightthickness=5)
            Contener.pack(expand=True,fill=BOTH)

            #List Frame of the Test
            Serology_Page = Frame(Contener)
            Serology_Page.pack(expand=1,fill=BOTH)
            Miscelaneous_Page = Frame(Contener)
            Urinalysis_Page= Frame(Contener)
            CBC_Page= Frame(Contener)
            FE_Page = Frame(Contener)

#Laboratory-END>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def Close_Plus_Finding(self):
        global PageOpen
        if messagebox.askokcancel('Close', 'Are you sure you want to close the Window all the data will not be Save?'):
            PageOpen=1
            self.Plus_Finding_Page.destroy()

    def Plus_Finding(self):
        global PageOpen
        if PageOpen < 2:
            self.Plus_Finding_Page=Toplevel()
            self.Plus_Finding_Page.title("Add Finding For X-Ray")
            Plus_Finding_width=600
            Plus_Finding_height=600
            self.Plus_Finding_Page.geometry(f'{Plus_Finding_width}x{Plus_Finding_height}+{400}+{80}')
            self.Plus_Finding_Page.protocol("WM_DELETE_WINDOW",self.Close_Plus_Finding)
            self.Plus_Finding_Page.resizable(False,False)

            Plus_Body=Frame(self.Plus_Finding_Page)
            Plus_Body.pack(expand=1,fill=BOTH)
            Plus_Title=Label(Plus_Body,text="ADD Finding",font=("Roboto",35,"bold")).place(x=5,y=10)

            Plus_Name=Label(Plus_Body,text="NAME:",font=("Roboto",12,"bold")).place(x=5,y=140)
            Plus_Name_Entry=Entry(Plus_Body,font=("Roboto",12),borderwidth=5)
            Plus_Name_Entry.place(x=65,y=137,relwidth=0.7)

            Label_Plus_Body=Label(Plus_Body,text="FINDING BODY:",font=("Roboto 15 bold"))
            Label_Plus_Body.place(x=5,y=170)
            Plus_Body_Text=Frame(Plus_Body,width=550,height=600,padx=5,pady=5,highlightbackground="black",highlightthickness=1)
            Plus_Body_Text.place(x=0,y=200)
            Plus__scroll=Scrollbar(Plus_Body_Text,orient='vertical')
            Plus__scroll.pack(side=RIGHT,fill='y')
            Plus__BOX = Text(Plus_Body_Text,width = 70,height = 19,borderwidth=5,font=("Roboto 11 "),yscrollcommand=Plus__scroll.set)
            Plus__scroll.config(command=Plus__BOX.yview)
            Plus__BOX.pack()

            def addFinding():
                global PageOpen
                self.user.addXrayFinding(Plus_Name_Entry.get(),Plus__BOX.get("1.0", "end-1c"))
                PageOpen = 1

                opts=self.user.getAllXrayFinding()
                n=2
                if opts is not None:
                    Option=[x[n] for x in opts]
                    self.FINDING_Mune.config(value=Option)
                self.Plus_Finding_Page.destroy()
                

            Plus_ADD_button=Button(Plus_Body,text="ADD",font=("Roboto 10"),width=5,borderwidth=5,command=addFinding)
            Plus_ADD_button.place(x=450,y=560)
           
            Plus_Cancel_button=Button(Plus_Body,text="Cancel",font=("Roboto 10"),width=5,borderwidth=5,command=lambda:self.Plus_Finding_Page.destroy())
            Plus_Cancel_button.place(x=520,y=560)

            PageOpen += 1
        else:
            messagebox.showinfo("Error","The Window is already Open!")

#X_Ray Laboratory>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def X_Ray(self):
        if self.user.return_dept()!='Imaging Center':
            messagebox.showerror("Access Denied","Only Employees from Imaging Center or Users with Administrative Access can Access this Page")
        else:
            self.Dashboard_GUI.resizable(True,True)
            width= self.Dashboard_GUI.winfo_screenwidth()
            height=self.Dashboard_GUI.winfo_screenheight()
            self.Dashboard_GUI.geometry(f'{width}x{height}+{0}+{0}')
            self.Page_Dashboard.forget()
            self.Page_XRAY=Frame(self.Dashboard_GUI)
            self.Page_XRAY.pack(expand=1, fill=BOTH)
            
            Frame_Header=Frame(self.Page_XRAY,bg='#BDFFC4',highlightbackground="black",highlightthickness=1)
            Frame_Header.pack(fill=X)
            image3 = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((40, 40)))
            IMG_HEADER_Xray=Label(Frame_Header,image=image3,bg='#BDFFC4',width=40,height=40)
            IMG_HEADER_Xray.image=image3
            IMG_HEADER_Xray.pack(side=LEFT,padx=20)
            HEADER_TITLE=Label(Frame_Header,text="City Health Office",bg='#BDFFC4',font='Roboto 20 bold')
            HEADER_TITLE.pack(side=LEFT)

            Toggle_Button=Menubutton(Frame_Header,width=5,text="=",highlightbackground="black",highlightthickness=1,justify=RIGHT)
            Toggle_Button.pack(side=RIGHT,padx=8)
            Toggle_Button.menu=Menu(Toggle_Button)
            Toggle_Button["menu"]=Toggle_Button.menu

            Toggle_Button.menu.add_command(label="Home",font='Roboto 12',command=lambda:self.HOME_PAGE(self.Close_ID[1]))
            Toggle_Button.menu.add_command(label="Logout",font='Roboto 12',command=lambda:self.logout())

            HEADER_USERNAME=Label(Frame_Header,text=str(self.user.username),bg='#BDFFC4',font='Roboto 20 ')
            HEADER_USERNAME.pack(side=RIGHT)
            #Header-------
            #BODY >> Laboratory
            Detail_Body=Frame(self.Page_XRAY)
            Detail_Body.pack(expand=1,fill=BOTH,side=LEFT)
            
            XRAY_Title=Label(Detail_Body,text="X-Ray Laboratory Test",anchor=W,font='Roboto 40 bold')
            XRAY_Title.pack(fill=X,padx=5,pady=20)

            Client_Detailbox=Frame(Detail_Body)
            Client_Detailbox.pack(pady=10)

            def setValue(event):
                global client_id, test_id
                res=self.user.getClient_name_Xray(Name_Entry.get())
                test_id=res[-1]
                self.Xray_Name_Entry.set(res[1])
                self.Birth_Entry.set_date(res[4])
                self.age.set(res[2])
                self.Xray_Gender_Mune.set(res[3])
                client_id=res[0]


            name=StringVar()
            Label(Client_Detailbox,text="Name: ",font='Roboto 12',anchor=W).pack(fill=X)
            self.Xray_Name_Entry=ttk.Combobox(Client_Detailbox,textvariable=name,font='Roboto 9',width=48)
            result=self.user.getClients_Xray()
            n=1
            self.Xray_Name_Entry['values']=[x[n] for x in result]
            self.Xray_Name_Entry.pack()
            self.Xray_Name_Entry.bind('<<ComboboxSelected>>', setValue)

            Birth_Label=Label(Client_Detailbox,text="Birthdate:",font="Roboto 12",anchor=W).pack(fill=X)
            self.Xray_Birth_Entry=DateEntry(Client_Detailbox,backgroud="magenta3",foreground="White",font="Roboto 12",bd=2,archor=W)
            self.Xray_Birth_Entry.pack(fill=X)

            def Name_entry_search(event):
                value=event.widget.get()
                if value!='':
                    lst=self.user.getClients_NameEntry_XraySearch(value)
                    data=[]
                    for item in range(len(lst)):
                        if value.lower() in lst[item][1].lower():
                            data.append(lst[item][1])
                    event.widget['values']=data

                else:
                    lst=self.user.getClients_Xray()
                    event.widget['values']=([x[1] for x in lst])

            self.Xray_Name_Entry.bind("<KeyRelease>",Name_entry_search)

            self.age=StringVar()
            AGE_Label=Label(Client_Detailbox,text="Age: ",font='Roboto 12',anchor=W).pack(fill=X)
            self.Xray_AGE_Entry=Entry(Client_Detailbox,textvariable=self.age,font='Roboto 9',borderwidth=3)
            self.Xray_AGE_Entry.pack(fill=X)

            def Gender_Click():
                Genderlabel=Label(Client_Detailbox,Gender_Mune.get(),font="Roboto 12 bold")

            Gender_Label=Label(Client_Detailbox,text="Gender:",font='Roboto 12 ',anchor=W).pack(fill=X)
            Option=["Male","Female","Other"]
            self.Xray_Gender_Mune=ttk.Combobox(Client_Detailbox,value=Option,font='Roboto 12',state='readonly')
            self.Xray_Gender_Mune.set("Select Gender")
            self.Xray_Gender_Mune.bind("<<ComboboxSelected>>",Gender_Click)
            self.Xray_Gender_Mune.pack(fill=X)

            def FIND_Click(event):
                chosen_finding_title=self.FINDING_Mune.get()
                findings=self.user.getXrayFindingDetails(chosen_finding_title)
                title=findings[0]
                body=findings[1]
                Finding_BOX.delete("1.0","end")
                # print(title,body)
                Finding_BOX.insert("1.0",body)

            opts=self.user.getAllXrayFinding()
            n=2
            if opts is not None:
                Option=[x[n] for x in opts]

            self.body=Frame(Detail_Body)
            self.body.pack()
            self.body2=Frame(self.body)
            self.body2.pack(fill=X)
            self.FINDING_Mune=ttk.Combobox(self.body2,value=Option,font='Roboto 12',width=20)
            self.FINDING_Mune.set("Select FINDING")
            self.FINDING_Mune.bind("<<ComboboxSelected>>",FIND_Click)
            self.FINDING_Mune.pack(side=RIGHT)
            Find_ADD=Button(self.body2,text="+",font='Roboto 10 bold',width=2,height=1,command=self.Plus_Finding)
            Find_ADD.pack(side=RIGHT)
            Label_Finding=Label(self.body2,text="Finding:",font=("Roboto 20 bold"))
            Label_Finding.pack(side=LEFT)
            FindBody=Frame(self.body,width=550,height=200,padx=5,pady=5,highlightbackground="black",highlightthickness=1)
            FindBody.pack()
            Find_scroll=Scrollbar(FindBody,orient='vertical')
            Find_scroll.pack(side=RIGHT,fill='y')
            Finding_BOX = Text(FindBody, height = 9, width = 70,borderwidth=5,font=("Roboto 11 "),yscrollcommand=Find_scroll.set)
            Find_scroll.config(command=Finding_BOX.yview)
            Finding_BOX.pack()

            self.body_IM=Frame(Detail_Body)
            self.body_IM.pack()
            Label_IMPRESSIONSBody=Label(self.body_IM,text="Impression:",font=("Roboto 20 bold"),anchor=W)
            Label_IMPRESSIONSBody.pack(fill=X)
            IMPRESSIONSBody=Frame(self.body_IM,width=550,height=200,padx=5,pady=5,highlightbackground="black",highlightthickness=1)
            IMPRESSIONSBody.pack()
            IMPRESSIONS_scroll=Scrollbar(IMPRESSIONSBody,orient='vertical')
            IMPRESSIONS_scroll.pack(side=RIGHT,fill='y')
            IMPRESSIONS_BOX = Text(IMPRESSIONSBody, height = 5, width = 70,borderwidth=5,font=("Roboto 11 "),yscrollcommand=IMPRESSIONS_scroll.set)
            IMPRESSIONS_scroll.config(command=IMPRESSIONS_BOX.yview)
            IMPRESSIONS_BOX.pack()

            Img_Body=Frame(self.Page_XRAY)
            Img_Body.pack(fill=BOTH,side=RIGHT,expand=True)

            Image_Box=Frame(Img_Body,width=300,height=390,bg="blue",highlightbackground="black",highlightthickness=2)
            Image_Box.pack(pady=40)

            Upload_button=Button(Image_Box,text="Upload the X-Ray",width=40,height=25,command=lambda:open_file())
            Upload_button.pack(fill=BOTH)

            def open_file():
                global img, filepath,b2
                f_types = [('Jpg Files', '*.jpg')]
                filepath = filedialog.askopenfilename(filetypes=f_types)
                img=Image.open(filepath)
                img_resized=img.resize((285,375))
                img=ImageTk.PhotoImage(img_resized)
                b2 =Button(Image_Box,image=img,borderwidth=5,command=Reopen) 
                b2.pack(fill=BOTH)
                Upload_button.pack_forget()
            
            def Reopen():
                if messagebox.askokcancel('Change', 'Are you sure you want to Delete this Image?'):
                    Upload_button.pack(fill=BOTH)
                    b2.pack_forget()

            def submit():
                if name.get() =="":
                    messagebox.showerror("Error","Select a Client First")
                else:
                    if Finding_BOX.get("1.0","end-1c")=="" and IMPRESSIONS_BOX.get("1.0","end-1c")=="" and 'filepath' not in globals():
                        messagebox.showerror("Error","Empty Finding, Impression and Xray Box")
                    elif Finding_BOX.get("1.0","end-1c")=="": messagebox.showerror("Error","Empty Finding Box")
                    elif IMPRESSIONS_BOX.get("1.0","end-1c")=="": messagebox.showerror("Error","Empty Impression Box")
                    elif filepath is None:messagebox.showerror("Error","Empty Xray Box")
                    else:
                        serviceid=self.user.get_test_id("X-Ray Test")
                        client_id=self.user.getClient_name_Xray(Name_Entry.get())
                        total=self.user.get_test_price(serviceid[0])
                        test_id=self.user.get_tests_id(client_id[0],serviceid[0])

                        document=Path(__file__).parent / os.getenv("XRAY_TEMPLATE")
                        doc=DocxTemplate(document)

                        image=InlineImage(doc,filepath,width=Inches(3), height=Inches(2.94))
                        context={
                            "NAME":name.get(),
                            "DOB":self.Xray_Birth_Entry.get_date(),
                            "AGE":age.get(),
                            "GENDER":Gender_Mune.get(),
                            "DATE":self.test_date,
                            "ID":client_id[0],
                            "FINDINGS":Finding_BOX.get("1.0","end-1c"),
                            "IMPRESSION":IMPRESSIONS_BOX.get("1.0","end-1c"),
                            "IMAGE":image,
                            "medtech_name":self.user.fname+" "+self.user.lname
                        }
                        doc.render(context)
                        path=os.getenv('DRIVE_PATH')+"/X-Ray"
                        Exists=os.path.exists(path)
                        if not Exists:
                            os.makedirs(path)

                        doc.save(path+"/"+name.get()+"_"+str(self.test_date)+".docx")
                        win32api.ShellExecute(0, "print", path+"/"+name.get()+"_"+str(self.test_date)+".docx", None, ".", 0)
                        
                        # self.user.markXray_as_done(test_id)

                        id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                        self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                        self.user.markTest_as_done(test_id[0])
            XRay_Continer=Frame(Img_Body)
            XRay_Continer.pack(side=BOTTOM,fill=X,pady=40)
            Submit_Xray=Button(XRay_Continer,text="Submit",width=10,bg="green",font='Roboto 11',borderwidth=5,command=lambda: submit()).pack(side=RIGHT,padx=20)
            Record_Xray=Button(XRay_Continer,text="Record",width=10,bg="green",font='Roboto 11',borderwidth=5,command=lambda:self.Record(self.Value_Laboratory[1])).pack(side=RIGHT)

#X_Ray Laboratory  END>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>

#Summary>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def Summary(self):
        width= self.Dashboard_GUI.winfo_screenwidth()
        height=self.Dashboard_GUI.winfo_screenheight()
        self.Dashboard_GUI.geometry(f'{width}x{height}+{0}+{0}')
        self.Dashboard_GUI.resizable(True,True)
        self.Page_Dashboard.forget()
        self.Page_Summary=Frame(self.Dashboard_GUI,bg="green")
        self.Page_Summary.pack(expand=1, fill=BOTH)
        Frame_Header=Frame(self.Page_Summary,bg='#BDFFC4',highlightbackground="black",highlightthickness=1)
        Frame_Header.pack(fill=X)
        image2 = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((40, 40)))
        IMG_HEADER_SUM=Label(Frame_Header,image=image2,bg='#BDFFC4',width=40,height=40)
        IMG_HEADER_SUM.image=image2
        IMG_HEADER_SUM.pack(side=LEFT)
        HEADER_TITLE=Label(Frame_Header,text="City Health Office",bg='#BDFFC4',font='Roboto 25 bold').pack(side=LEFT)

        Toggle_Button=Menubutton(Frame_Header,width=5,text="=",highlightbackground="black",highlightthickness=1,justify=RIGHT)
        Toggle_Button.pack(side=RIGHT,padx=20)
        Toggle_Button.menu=Menu(Toggle_Button)
        Toggle_Button["menu"]=Toggle_Button.menu
        Toggle_Button.menu.add_command(label="Home",font='Roboto 12 ',command=lambda:self.HOME_PAGE(self.Close_ID[2]))
        Toggle_Button.menu.add_command(label="Logout",font='Roboto 12 ',command=lambda:self.logout())

        HEADER_USERNAME=Label(Frame_Header,text=str(self.user.username),font='Roboto 20 ',bg='#BDFFC4')
        HEADER_USERNAME.pack(side=RIGHT)
        #Header-------
        #BODY >> Summary

        Frame_SumBody=Frame(self.Page_Summary)
        Frame_SumBody.pack(expand=1,fill=BOTH,side=TOP)
        
        Frame_FilterBody=Frame(Frame_SumBody,height=130,border=2,borderwidth=5,highlightbackground="black",highlightthickness=1)
        Frame_FilterBody.pack(fill=X)

        FrameIMG1=Frame(Frame_FilterBody, bg="blue",width=110,height=110)
        FrameIMG1.pack(side=LEFT,padx=10)

        image = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((110, 110)))
        Img1=Label(FrameIMG1,image = image)
        Img1.image=image
        Img1.pack()

        res= self.user.getAllTest()
        Sum_Test=[x[0] for x in res]
        Sum_Test.insert(0,'All')

        Container=Frame(Frame_FilterBody)
        Container.pack(side=LEFT,fill=Y,padx=10)

        Label(Container,text="Laboratory Test:",font='Roboto 11',anchor=W).pack(fill=X)
        LabTest_Test=ttk.Combobox(Container,value=Sum_Test,font='Roboto 10',state='readonly')
        LabTest_Test.set("All")
        LabTest_Test.pack(fill=X)

        emp=employee.Employee.getAllEmployees()
        emp_choices=[x[1]+' '+x[2] for x in emp]
        emp_choices.insert(0,'All')
        # Label(Container,text="").pack(pady=1)
        Label(Container,text="Medical Technologist",font='Roboto 11',anchor=W).pack(fill=X)
        MidTech_Emp=ttk.Combobox(Container,value=emp_choices,font='Roboto 10',state='readonly',width=45)
        MidTech_Emp.set("All")
        MidTech_Emp.pack()

        Container_2=Frame(Frame_FilterBody)
        Container_2.pack(side=LEFT,fill=Y,padx=5)

        filter_options=["No Filter","Monthly","Yearly","1st Semi Annual","2nd Semi Annual","1st Quarter","2nd Quarter","3rd Quarter","4th Quarter"]
        Label(Container_2,text="Filter By:",font='Roboto 11',anchor=W).pack(fill=X)
        MidTech_Filter=ttk.Combobox(Container_2,value=filter_options,font='Roboto 10',state='readonly',width=20)
        MidTech_Filter.set("No Filter")
        MidTech_Filter.pack()

        global Valuebox, Monthly_year

        Valuebox_V=list(calendar.month_name)
        Valuebox_label=Label(Container_2,text="Choose Month:",font='Roboto 11',anchor=W)

        Valuebox=ttk.Combobox(Container_2,value=Valuebox_V,font='Roboto 10',state='readonly',width=20)
        Valuebox.set("Select Month")

        Monthly_yearr=datetime.today().year
        Monthly_years=[Monthly_yearr - i for i in range (6)]
        Monthly_Lyears=Label(Container_2,text="Choose Year :",font='Roboto 11',anchor=W)
        Monthly_year=ttk.Combobox(Container_2,value=Monthly_years,font='Roboto 10',state='readonly',width=20)
        Monthly_year.set("Select Year")


        #Time to View and Save Graph undone
        '''Graph_Label=Label(Frame_FilterBody,text="Graph Model:",font='Roboto 11').place(x=850,y=14)
        Graph_Value=["GRAPH BAR","GRAPH PIE"]        
        Graph_Selection=ttk.Combobox(Frame_FilterBody,value=Graph_Value,font='Roboto 10',state='readonly',width=20)
        Graph_Selection.set("Select Graph Model")
        Graph_Selection.place(x=940,y=14)

        def Graph_create():
            Selected_value = Graph_Selection.get()
            if Selected_value== "GRAPH BAR":
                fig, ax = plt.subplots()
                fruits = ['apple', 'blueberry', 'cherry', 'orange']
                counts = [40, 100, 30, 55]
                bar_labels = ['red', 'blue', '_red', 'orange']
                bar_colors = ['tab:red', 'tab:blue', 'tab:red', 'tab:orange']
                ax.bar(fruits, counts, label=bar_labels, color=bar_colors)
                ax.set_ylabel('fruit supply')
                ax.set_title('Fruit supply by kind and color')
                ax.legend(title='Fruit color')

                plt.show()

            elif Selected_value== "GRAPH PIE":
                labels = 'Frogs', 'Hogs', 'Dogs', 'Logs'
                sizes = [15, 30, 45, 10]
                explode = (0, 0.1, 0, 0)  # only "explode" the 2nd slice (i.e. 'Hogs')
                fig1, ax1 = plt.subplots()
                ax1.pie(sizes, explode=explode, labels=labels, autopct='%1.1f%%',
                        shadow=True, startangle=90)
                ax1.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
                plt.show()
            
            else:
                messagebox.showerror("Access Denied","Please Select A Graph Model to Use For Data Analyst")

        Graph_button = Button(Frame_FilterBody,text="Create Graph",font='Roboto 7',borderwidth=3,command=Graph_create)
        Graph_button.place(x=1100,y=12)'''

        global name_Choice,test_choice, filter_choice,filter_date_from,filter_date_to

        name_Choice=None
        test_choice=None
        filter_choice=None
        filter_date_from=None
        filter_date_to=None

        def filter_Option(event):
            global filter_choice
            filter_choice=event.widget.get()
            if event.widget.get() == "No Filter":
                filter_choice=None
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.pack_forget()
                Valuebox.pack_forget()


            elif event.widget.get()=="Monthly":
                Monthly_months=list(calendar.month_name)
                Valuebox_label.config(text="Choose Month:")
                Valuebox.config(values=Monthly_months)

                Valuebox_label.pack(fill=X)
                Valuebox.pack()

                Monthly_Lyears.pack(fill=X)
                Monthly_year.pack()

            elif event.widget.get()=="Yearly":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="Choose Yearly:")
                Valuebox_label.pack(fill=X)
                Valuebox.pack()

                Yearly_yearr=datetime.today().year
                Yearly_years=[Yearly_yearr - i for i in range (6)]
                Valuebox.config(values=Yearly_years)
                Valuebox.set('Select Year')

            elif event.widget.get()=="1st Semi Annual":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="1st Semi Annual:")
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Semi_1=datetime.today().year
                Semi_1_years=[Semi_1 - i for i in range (6)]
                Valuebox.config(values=Semi_1_years)
                Valuebox.set('Select Year')

            elif event.widget.get()=="2nd Semi Annual":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="2nd Semi Annual:")
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Semi_2=datetime.today().year
                Semi_2_years=[Semi_2 - i for i in range (6)]
                Valuebox.config(values=Semi_2_years)
                Valuebox.set('Select Year')
            
            elif event.widget.get()=="1st Quarter":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="1st Quarter:")
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Quarter_1=datetime.today().year
                Quarter_1_years=[Quarter_1 - i for i in range (6)]
                Valuebox.config(values=Quarter_1_years)
                Valuebox.set('Select Year')

            elif event.widget.get()=="2nd Quarter":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="2nd Quarter:")
                Valuebox.set('Select Year')
                
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Quarter_2=datetime.today().year
                Quarter_2_years=[Quarter_2 - i for i in range (6)]
                Valuebox.config(values=Quarter_2_years)
                Valuebox.set('Select Year')

            elif event.widget.get()=="3rd Quarter":
                Valuebox_label.config(text="3rd Quarter:")
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Quarter_3=datetime.today().year
                Quarter_3_years=[Quarter_3 - i for i in range (6)]
                Valuebox.config(values=Quarter_3_years)
                Valuebox.set('Select Year')
                
            elif event.widget.get()=="4th Quarter":
                Monthly_Lyears.pack_forget()
                Monthly_year.pack_forget()
                Valuebox_label.config(text="4th Quarter:")
                Valuebox_label.pack(fill=X)
                Valuebox.pack()
                Quarter_4=datetime.today().year
                Quarter_4_years=[Quarter_4 - i for i in range (6)]
                Valuebox.config(values=Quarter_4_years)
                Valuebox.set('Select Year')
                
        MidTech_Filter.bind("<<ComboboxSelected>>",filter_Option)

        Frame_TableBody=Frame(Frame_SumBody,bg="grey",borderwidth=5,highlightbackground="black",highlightthickness=1)
        Frame_TableBody.pack(expand=1,fill=BOTH)
        Summary_Table=ttk.Treeview(Frame_TableBody)
        style=ttk.Style()
        style.theme_use("default")
        style.configure("Treeview")
        Summary_Table['column']=("ID","NAME","GENDER","TEST","AGE","DATE STARTED","DATE FINISHED","MEDTECH")
        #Column
        Summary_Table.column("#0",width=0,stretch=NO)
        Summary_Table.column("ID",width=50,stretch=NO,anchor=CENTER)
        Summary_Table.column("NAME")
        Summary_Table.column("GENDER")
        Summary_Table.column("TEST",width=200,stretch=NO)
        Summary_Table.column("AGE")
        Summary_Table.column("DATE STARTED",width=100,stretch=NO,anchor=CENTER)
        Summary_Table.column("DATE FINISHED",width=100,stretch=NO,anchor=CENTER)
        Summary_Table.column("MEDTECH")
        #Header
        Summary_Table.heading("#0")
        Summary_Table.heading("ID",text="No")
        Summary_Table.heading("NAME",text="NAME",anchor=W)
        Summary_Table.heading("GENDER",text="Gender",anchor=W)
        Summary_Table.heading("TEST",text="TEST")
        Summary_Table.heading("AGE",text="Age",anchor=W)
        Summary_Table.heading("DATE STARTED",text="DATE STARTED",anchor=W)
        Summary_Table.heading("DATE FINISHED",text="DATE FINISHED",anchor=W)
        Summary_Table.heading("MEDTECH",text="Medical Technologists",anchor=W)
        Summaryscroll=ttk.Scrollbar(Frame_TableBody,orient=VERTICAL,command=Summary_Table.yview)
        Summaryscroll.pack(side=RIGHT,fill=Y)
        Summary_Table.pack(expand=1,fill=BOTH)

        res=self.user.getAllClient_Done()
        count=0
        number=1
        for item in range(len(res)):
            Summary_Table.insert(parent='',index='end',iid=count,value=(number,res[item][1],res[item][2],res[item][3],res[item][4],res[item][5],res[item][6],res[item][7]))
            count+=1
            number+=1


        def LabTest_callback(event):
            global test_choice
            test_choice=event.widget.get()
            
        def nameCallback(event):
            global name_Choice
            name_Choice=event.widget.get()

        LabTest_Test.bind("<<ComboboxSelected>>",LabTest_callback)
        MidTech_Emp.bind("<<ComboboxSelected>>",nameCallback)

        def ApplyFilter():
            months=['January','February',"March","April","May","June","July","August","September","October","November","December"]
            global name_Choice,test_choice,filter_date_from,filter_date_to,filter_choice
            if filter_choice=="Monthly":
                if Valuebox.get()=="Select Month" or Monthly_year.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Month and/or Year is Empty.")
                elif Valuebox.get()!="Select Month" and Valuebox.get() in months:
                    if Valuebox.get()=="January":
                        month_num=1
                    elif Valuebox.get()=="February":
                        month_num=2
                    elif Valuebox.get()=="March":
                        month_num=3
                    elif Valuebox.get()=="April":
                        month_num=4
                    elif Valuebox.get()=="May":
                        month_num=5
                    elif Valuebox.get()=="June":
                        month_num=6
                    elif Valuebox.get()=="July":
                        month_num=7
                    elif Valuebox.get()=="August":
                        month_num=8
                    elif Valuebox.get()=="September":
                        month_num=9
                    elif Valuebox.get()=="October":
                        month_num=10
                    elif Valuebox.get()=="November":
                        month_num=11
                    elif Valuebox.get()=="December":
                        month_num=12

                    filter_date_from_dateObj=date(int(Monthly_year.get()),month_num,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(filter_date_from_dateObj.year, filter_date_from_dateObj.month)
                    day=res[1]
                    filter_date_to=date(int(Monthly_year.get()),month_num,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="Yearly":
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()
            elif filter_choice=="1st Semi Annual": 
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),1,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 6)
                    day=res[1]
                    filter_date_to=date(int(year),6,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="2nd Semi Annual": 
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),7,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 12)
                    day=res[1]
                    filter_date_to=date(int(year),12,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="1st Quarter":
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),1,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 3)
                    day=res[1]
                    filter_date_to=date(int(year),3,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="2nd Quarter":
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),4,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 6)
                    day=res[1]
                    filter_date_to=date(int(year),6,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="3rd Quarter":
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),7,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 9)
                    day=res[1]
                    filter_date_to=date(int(year),9,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            elif filter_choice=="4th Quarter":
                if Valuebox.get()=="Select Year":
                    messagebox.showerror("No Value Selected", "Year is Empty.")
                elif Valuebox.get()!="Select Year":
                    year=Valuebox.get()

                    filter_date_from_dateObj=date(int(year),10,1)
                    filter_date_from=filter_date_from_dateObj.strftime("%Y-%m-%d")

                    res=calendar.monthrange(int(year), 12)
                    day=res[1]
                    filter_date_to=date(int(year),12,day)
                    filter_date_to=filter_date_to.strftime("%Y-%m-%d")

            if name_Choice is None:
                name_Choice="All"
            if test_choice is None:
                test_choice="All"
            sum=summary_filter.Summary()
            if filter_choice is None and name_Choice!="All" and test_choice!="All" or  filter_choice is None and  name_Choice=="All" and test_choice=="All" or  filter_choice is None and name_Choice=="All" and test_choice!="All" or  filter_choice is None and name_Choice!="All" and test_choice=="All":
                res=sum.filterOut(name_Choice,test_choice)

            elif filter_choice=="Monthly" and name_Choice=="All" and test_choice=="All":
                res=sum.filterMonthly(filter_date_from,filter_date_to)

            elif filter_choice=="Monthly" and name_Choice!="All" and test_choice=="All" or filter_choice=="Monthly" and name_Choice=="All" and test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice)

            elif filter_choice=="Yearly" and name_Choice!="All" and test_choice=="All" or filter_choice=="Yearly" and name_Choice!="All" and test_choice!="All" or filter_choice=="Yearly" and name_Choice=="All" and test_choice!="All":

                res=sum.filterYearlyTest(year,name_Choice,test_choice)
            elif filter_choice=="Yearly" and name_Choice=="All" and test_choice=="All":

                res=sum.filterYearly(year)

            elif filter_choice=='1st Semi Annual' and name_Choice=="All" and test_choice=="All" or filter_choice=='2nd Semi Annual' and name_Choice=="All" and test_choice=="All":

                res=sum.filterMonthly(filter_date_from,filter_date_to)
            elif filter_choice == '2nd Semi Annual' and name_Choice!="All" or test_choice!="All" or filter_choice == '2nd Semi Annual' and name_Choice!="All" or test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice) 

            elif filter_choice=='1st Quarter' and name_Choice=="All" and test_choice=="All" or filter_choice=='1st Quarter' and name_Choice=="All" and test_choice=="All":
                res=sum.filterMonthly(filter_date_from,filter_date_to)

            elif filter_choice == '1st Quarter' and name_Choice!="All" or test_choice!="All" or filter_choice == '1st Quarter' and name_Choice!="All" or test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice)  

            elif filter_choice=='2nd Quarter' and name_Choice=="All" and test_choice=="All" or filter_choice=='2nd Quarter' and name_Choice=="All" and test_choice=="All":

                res=sum.filterMonthly(filter_date_from,filter_date_to)
            elif filter_choice == '2nd Quarter' and name_Choice!="All" or test_choice!="All" or filter_choice == '2nd Quarter' and name_Choice!="All" or test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice)  

            elif filter_choice=='3rd Quarter' and name_Choice=="All" and test_choice=="All" or filter_choice=='3rd Quarter' and name_Choice=="All" and test_choice=="All":

                res=sum.filterMonthly(filter_date_from,filter_date_to)
            elif filter_choice == '3rd Quarter' and name_Choice!="All" or test_choice!="All" or filter_choice == '3rd Quarter' and name_Choice!="All" or test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice)  

            elif filter_choice=='4th Quarter' and name_Choice=="All" and test_choice=="All" or filter_choice=='4th Quarter' and name_Choice=="All" and test_choice=="All":
                res=sum.filterMonthly(filter_date_from,filter_date_to)

            elif filter_choice == '4th Quarter' and name_Choice!="All" or test_choice!="All" or filter_choice == '4th Quarter' and name_Choice!="All" or test_choice!="All":
                res=sum.filterMonthlyTest(filter_date_from,filter_date_to,name_Choice,test_choice) 
 
            count=0
            number=1
            Summary_Table.delete(*Summary_Table.get_children())
            for item in range(len(res)):
                Summary_Table.insert(parent='',index='end',iid=count,value=(number,res[item][1],res[item][2],res[item][3],res[item][4],res[item][5],res[item][6],res[item][7]))
                count+=1
                number+=1
                
        def PrintResults():
            all_items=Summary_Table.get_children()
            tests={}
            Gender={}
            age={}
            testXMale={}
            testXFemale={}
            testXOther={}
            TestXAge={}

            for item in Summary_Table.get_children():
                itemValues=Summary_Table.item(item)['values']
                
                if itemValues[3] not in tests.keys():
                    tests[itemValues[3]]=1
                elif itemValues[3] in tests.keys():
                    tests[itemValues[3]]+=1

                if itemValues[2] not in Gender.keys():
                    Gender[itemValues[2]]=1
                else:
                    Gender[itemValues[2]]+=1

                if itemValues[4] not in age.keys():
                    age[itemValues[4]]=1
                else:
                    age[itemValues[4]]+=1

                if itemValues[3] not in testXMale.keys():
                    if itemValues[2]=="Male":
                        testXMale[itemValues[3]]=1
                else:
                    testXMale[itemValues[3]]+=1
                
                if itemValues[3] not in testXFemale.keys():
                    if itemValues[2]=="Female":
                        testXFemale[itemValues[3]]=1
                else:
                    testXFemale[itemValues[3]]+=1

                if itemValues[3] not in testXOther.keys():
                    if itemValues[2]=="Other":
                        testXOther[itemValues[3]]=1 
                else:
                    testXOther[itemValues[3]]+=1

                # TestXAge[1]=itemValues[0]
                # TestXAge['test']=itemValues[3]
                # TestXAge['age']=itemValues[4]

            fig, ax = plt.subplots()
            if LabTest_Test.get()=="All":
                bars=ax.barh(list(tests.keys()), list(tests.values()), label=tests.keys())
                plt.ylabel("Number of Clients per Test")
                # plt.legend(handles=bars, bbox_to_anchor=(0, 1.02), loc='lower left')
                plt.savefig('tests.png', dpi=300, bbox_inches='tight')

                fig, ax = plt.subplots()
                ax.pie(Gender.values(), labels=Gender.keys(), autopct='%1.0f%%')
                plt.savefig('gender.png', dpi=300)

                fig, ax = plt.subplots()
                ax.pie(age.values(), labels=age.keys(), autopct='%1.0f%%')
                plt.savefig('age.png', dpi=300)

                fig, ax = plt.subplots()
                bars=ax.barh(list(testXMale.keys()), list(testXMale.values()), label=testXMale.keys())
                plt.ylabel("Number of Males")
                # plt.legend(handles=bars, bbox_to_anchor=(0, 1.02), loc='lower left')
                plt.savefig('testXMale.png', dpi=300, bbox_inches='tight')

                fig, ax = plt.subplots()
                bars=ax.barh(list(testXFemale.keys()), list(testXFemale.values()), label=testXFemale.keys())
                plt.ylabel("Number of Females")
                # plt.legend(handles=bars, bbox_to_anchor=(0, 1.02), loc='lower left')
                plt.savefig('testXFemale.png', dpi=300, bbox_inches='tight')

                fig, ax = plt.subplots()
                bars=ax.barh(list(testXOther.keys()), list(testXOther.values()), label=testXOther.keys())
                plt.ylabel("Number of 'Other'")
                # plt.legend(handles=bars, bbox_to_anchor=(0, 1.02), loc='lower left')
                plt.savefig('testXOther.png', dpi=300, bbox_inches='tight')
            else:
                fig, ax = plt.subplots()
                ax.pie(Gender.values(), labels=Gender.keys(), autopct='%1.0f%%')
                plt.savefig('gender.png', dpi=300)

                fig, ax = plt.subplots()
                ax.pie(age.values(), labels=age.keys(), autopct='%1.0f%%')
                plt.savefig('age.png', dpi=300)
            

            doc = docx.Document(os.getenv("SUMMARY_REPORT_TEMPLATE"))
            data=[]
            for i in all_items:
                itemVal=Summary_Table.item(i, "values")
                data.append(itemVal)

            menuTable = doc.add_table(rows=1,cols=8)
            hdr_Cells = menuTable.rows[0].cells
            hdr_Cells[0].text="No"
            hdr_Cells[1].text="Name"
            hdr_Cells[2].text="Gender"
            hdr_Cells[3].text="Test"
            hdr_Cells[4].text="Age"
            hdr_Cells[5].text="Date Started"
            hdr_Cells[6].text="Date Finished"
            hdr_Cells[7].text="Medical Technologist"

            for no, name,gender,test,age, dateStart,dateFin,medTech in data:
                row_cell=menuTable.add_row().cells
                row_cell[0].text=str(no)
                row_cell[1].text=name
                row_cell[2].text=gender
                row_cell[3].text=test
                row_cell[4].text=age
                row_cell[5].text=dateStart
                row_cell[6].text=dateFin
                row_cell[7].text=medTech

            doc.add_picture('tests.png')
            doc.add_picture('gender.png')
            doc.add_picture('age.png')
            doc.add_picture('testXMale.png')
            doc.add_picture('testXFemale.png')
            doc.add_picture('testXOther.png')

            doc.save("new_document.docx")
            win32api.ShellExecute(0, "print", str(Path(__file__).parent/"new_document.docx"), None, ".", 0)

        Button(Frame_FilterBody,text="Apply Filter",command=lambda: ApplyFilter()).place(x=1000,y=50)
        Button(Frame_FilterBody,text="Print Out Results",command=lambda: PrintResults()).place(x=1100,y=50)

    def onClose(self):
        # windll.user32.ShowWindow(h, 9)
        self.Dashboard_GUI.destroy()

    def Cer_onClose(self):
        global PageOpen
        if messagebox.askokcancel('Close', 'Are you sure you want to close the View Page all the data will not be Save?'):
            PageOpen=1
            self.Dashboard_GUI.grab_release()
            self.Certificate.destroy()

    def Certificate_Page(self):
        global PageOpen
        if PageOpen < 2:
            self.Certificate = Toplevel(self.Dashboard_GUI)
            self.Certificate.title("Medical Certificate")
            Record_width=400
            Record_height=400
            self.Certificate.geometry(f'{Record_width}x{Record_height}+{480}+{100}')
            self.Certificate.resizable(False,False)
            self.Certificate.protocol("WM_DELETE_WINDOW", self.Cer_onClose)
            self.Certificate.grab_set()

            Certificate_title=Label(self.Certificate,text="Medical Certificate",font='Roboto 30 bold').place(x=7,y=7)
            res=self.user.getClientsMedCert()
            Patent_Label=Label(self.Certificate,text="Patients",font='Roboto 11').place(x=48,y=130)
            Patent_Value=[x[1] for x in res]        
            Patent_Selection=ttk.Combobox(self.Certificate,value=Patent_Value,font='Roboto 10',state='readonly',width=40)
            Patent_Selection.set("Select Patients")
            Patent_Selection.place(x=50,y=150)

            def setClient(event):
                res=self.user.getClient_name(Patent_Selection.get())
                global name, age, address
                name = res[1]
                age = res[2]
                address = res[5]

            Patent_Selection.bind("<<ComboboxSelected>>",setClient)

            P_lABEL=Label(self.Certificate,text="Please Select the Patients, its Purpose \nand its Validity",font='Roboto 13').place(x=50,y=70)

            PURPOSE_lABEL=Label(self.Certificate,text="Purpose",font='Roboto 11').place(x=48,y=180)
            PURPOSE_ENTRY=Entry(self.Certificate,width=33,borderwidth=3,font='Roboto 12')
            PURPOSE_ENTRY.place(x=50,y=200)

            REMARKS_lABEL=Label(self.Certificate,text="Remarks",font='Roboto 11').place(x=48,y=230)
            REMARKS_ENTRY=Entry(self.Certificate,width=33,borderwidth=3,font='Roboto 12')
            REMARKS_ENTRY.place(x=50,y=250)

            VALID_FROM_LABEL=Label(self.Certificate,text="Valid Until:",font='Roboto 11').place(x=48,y=280)
            VALID_FROM=DateEntry(self.Certificate,width=10,backgroud="magenta3",foreground="White",font="Roboto 12",bd=2,state='readonly')
            VALID_FROM.place(x=50, y=300)



            def submit():
                serviceid=self.user.get_test_id("Medical Certificate")
                client_id=self.user.getClient_name(name)
                amount=self.user.getTestAmount(serviceid)

                document=Path(__file__).parent / os.getenv("MED_CERT_TEMPLATE")
                doc=DocxTemplate(document)
                OR_Num=self.user.generateClient_ORNumber(),    
                context={
                    "DATE_TODAY":self.test_date,
                    "CLIENT_NAME":name,
                    "AGE":age,
                    "CLIENT_ADDRESS":address,
                    "PURPOSE": PURPOSE_ENTRY.get(),
                    "REMARKS": REMARKS_ENTRY.get(),
                    "VALID_UNTIL": str(self.test_date)+'-'+str(VALID_FROM.get_date()),
                    "OR_NUM": OR_Num[0],
                    "AMOUNT": amount[0],
        
                    "NAME_OF_DOCTOR":self.user.fname+" "+self.user.lname,
                    "POSITION":self.user.role,
                    "LICENSE_NO": self.user.license_no
                }
                doc.render(context)
                path=os.getenv('DRIVE_PATH')+"/Medical Certificates"
                Exists=os.path.exists(path)
                if not Exists:
                    os.makedirs(path)

                doc.save(path+"/"+name+"_"+str(OR_Num[0])+".docx")
                win32api.ShellExecute(0, "print", path+"/"+name+"_"+str(OR_Num[0])+".docx", None, ".", 0)

                total=self.user.get_test_price(serviceid[0])
                id=self.user.save_to_summary(total[0],serviceid[0],client_id[0])
                self.user.update_summaryID_test(id,client_id[0],serviceid[0])
                test_id=self.user.get_tests_id(client_id[0],serviceid[0])
                self.user.markTest_as_done(test_id[0][0])

            Certificate_button=Button(self.Certificate,text="Print Certificate",width=12,height=1,bg="green",borderwidth=5,command=submit).place(x=260,y=350)

            PageOpen += 1

        else:
            messagebox.showinfo("Error","The Window is already Open!")
            
    #Dashboard>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>
    def Main_Dashboard(self):   
        self.Dashboard_GUI=Tk()
        self.Dashboard_GUI.title('CITY HEALTH')
        # width= self.Dashboard_GUI.winfo_screenwidth()
        # height=self.Dashboard_GUI.winfo_screenheight()
        width= 1000
        height= 500
        self.Dashboard_GUI.geometry(f'{width}x{height}+{180}+{80}')
        self.Dashboard_GUI.protocol("WM_DELETE_WINDOW", self.onClose)
        self.Dashboard_GUI.resizable(False,False)

        self.Page_Dashboard=Frame(self.Dashboard_GUI)
        self.Page_Dashboard.pack(expand=1, fill=BOTH)
        Frame_Header=Frame(self.Page_Dashboard,bg='#BDFFC4',highlightbackground="black",highlightthickness=1)
        Frame_Header.pack(fill=X)
        image = ImageTk.PhotoImage(Image.open("CHO_LOGO.png").resize((40, 40)))
        IMG_HEADER_MD=Label(Frame_Header,image=image,bg='#BDFFC4',width=40,height=40)
        IMG_HEADER_MD.pack(side=LEFT)
        HEADER_TITLE=Label(Frame_Header,text="City Health Office",bg='#BDFFC4',font='Roboto 25 bold').pack(side=LEFT)

        Toggle_Button=Menubutton(Frame_Header,width=5,text="=",bg="white",highlightbackground="black",highlightthickness=1,justify=RIGHT)
        Toggle_Button.pack(side=RIGHT,padx=20)
        Toggle_Button.menu=Menu(Toggle_Button)
        Toggle_Button["menu"]=Toggle_Button.menu
        Toggle_Button.menu.add_command(label="LOGOUT",font='Roboto 12 bold',command=lambda:self.logout())

        HEADER_USERNAME=Label(Frame_Header,text=str(self.user.username),bg='#BDFFC4',font='Roboto 20')
        HEADER_USERNAME.pack(side=RIGHT)
        #Header END------------

        #Message
        Frame_Center=Frame(self.Page_Dashboard,highlightbackground="black",highlightthickness=1)
        Frame_Center.pack(expand=1,fill=BOTH)

        #FrontDesk
        Frame_Laboratory=Frame(self.Page_Dashboard,highlightbackground="black",highlightthickness=1)
        Frame_Laboratory.pack(fill=X)

        Frame_Profile = Frame(Frame_Laboratory,highlightbackground="black",highlightthickness=1)
        Frame_Profile.pack(expand=True,side=LEFT,fill=X)

        CHO_F= ImageTk.PhotoImage(Image.open("CHO_Front.jpg").resize((140, 140)))
        CHO_FL=Label(Frame_Profile,image=CHO_F,highlightbackground="black",highlightthickness=1)
        CHO_FL.pack(side=LEFT)


        USER_ID=Label(Frame_Profile,text="ID: "+str(self.user.id),font='Roboto 13')
        USER_ID.place(x=150,y=5)
        USER_NAME=Label(Frame_Profile,text="Name: "+str(self.user.username),font='Roboto 15')
        USER_NAME.place(x=150,y=30)
        USER_Role=Label(Frame_Profile,text="Profession:"+str(self.user.role),font='Roboto 13')
        USER_Role.place(x=150,y=60)
        # USER_NAME=Label(Frame_Profile,text=str(self.user.username),font='Roboto 20')
        # USER_NAME.pack(side=LEFT)

        Frame_Works = Frame(Frame_Laboratory,width=200,highlightbackground="black",highlightthickness=1)
        Frame_Works.pack(fill=BOTH,side=RIGHT)

        if self.user.return_dept()=='Imaging Center':
            Button_XRay=Button(Frame_Works,text="X_RAY",font=("Roboto",12,"bold"),width=10,height=2,bg="green",borderwidth=5,command=self.X_Ray)
            Button_XRay.pack(pady=2)
            Certificate_Button=Button(Frame_Works,text="Certificate",font=("Roboto",12,"bold"),width=10,height=1,bg="green",borderwidth=5,command=self.Certificate_Page)
            Certificate_Button.pack(pady=1)

        elif self.user.return_dept()=='Laboratory Department':
            Button_LabCH=Button(Frame_Works,text="Laboratory",font=("Roboto",12,"bold"),width=10,height=2,bg="green",borderwidth=5,command=self.Laboratory)
            Button_LabCH.pack(pady=2)
            Certificate_Button=Button(Frame_Works,text="Certificate",font=("Roboto",12,"bold"),width=10,height=1,bg="green",borderwidth=5,command=self.Certificate_Page)
            Certificate_Button.pack(pady=1)

        else:
            # messagebox.showinfo("Error","This Account don't have a Role!")
            Button_FronDesk=Button(Frame_Works,text="FrontDisk",font=("Roboto",12,"bold"),width=10,height=3,bg="green",borderwidth=5,command=self.FrontDesk)
            Button_FronDesk.pack()

        Button_Summary=Button(Frame_Works,text="Summary",font=("Roboto",12,"bold"),width=10,height=1,bg="green",borderwidth=5,command=self.Summary)
        Button_Summary.pack()
        
        self.Dashboard_GUI.mainloop()
    
    def start(self):
        self.Main_Dashboard()